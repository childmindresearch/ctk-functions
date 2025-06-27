"""Integration tests (with LanguageTool) for document corrections."""

import aiohttp
import docx
import pytest
import pytest_mock

from ctk_functions.routers.intake.intake_processing.utils import language_utils


@pytest.mark.asyncio
async def test_document_corrections(mocker: pytest_mock.MockerFixture) -> None:
    """Tests the DocumentCorrections class happy path.

    These runs are directly copied from a real paragraph in the test document, with
    the name/pronouns filled in as Lea she/her.
    """
    session_spy = mocker.patch(
        "ctk_functions.microservices.language_tool.aiohttp.ClientSession",
        wraps=aiohttp.ClientSession,
    )

    document = docx.Document()
    run_texts = [
        "At the standardized testing sessions,",
        " ",
        "Lea",
        " ",
        "presented as a (",
        "e.g.",
        "cooperative, friendly, hard - working, playful, gregarious, mature",
        "etc.)",
        " ",
        "girl",
        " ",
        "and appeared",
        " her",
        " s",
        "tated",
        " age.",
        " she",
        " was",
        " casually / formally",
        " dressed",
        " with appropriate grooming and hygiene.",
        " ",
        "she",
        " ",
        "displayed good eye contact and",
        " ",
        "social reciprocity.",
        " ",
        "she",
        " ",
        "was appropriately talkative, and",
        " ",
        "her",
        " ",
        "speech was normal i",
        "n pace, rate, and volume.",
        " ",
        "Lea",
        " ",
        (
            "presented as euthymic / dysthymic with full range of affective expression,"
            " which appeared congruent with the situation."
        ),
        " ",
        "Her",
        " ",
        (
            "thought process was logical / circumstantial / goal-directed, and no "
            "flight of ideas or loose associations were evident."
        ),
        " ",
    ]
    full_text = "".join(run_texts)
    paragraph = document.add_paragraph()
    for text in run_texts:
        paragraph.add_run(text)
    correcter = language_utils.DocumentCorrections(document)

    await correcter.correct()
    assert session_spy.call_count == 1

    assert paragraph.text == full_text.replace("she", "She")
