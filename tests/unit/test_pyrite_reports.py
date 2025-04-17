"""Tests for the reports construction."""

import cmi_docx
import docx

from ctk_functions.core import config
from ctk_functions.routers.pyrite import reports
from ctk_functions.routers.pyrite.tables import scq

settings = config.get_settings()


def test_run_section_add_to() -> None:
    """Test adding a RunsSection to a document."""
    doc = docx.Document()
    section = reports.RunsSection(
        content=(
            "Test ",
            "content",
        ),
        run_styles=("Emphasis", None),
    )

    section.add_to(doc)

    assert doc.paragraphs[0].runs[0].text == "Test "
    assert doc.paragraphs[0].runs[1].text == "content"
    assert doc.paragraphs[0].runs[0].style.name == "Emphasis"
    assert doc.paragraphs[0].runs[1].style.name == "Default Paragraph Font"


def test_paragraph_section_add_to() -> None:
    """Test adding a ParagraphSection to a document."""
    doc = docx.Document()
    section = reports.ParagraphSection(content="Hello world!", style="Heading 1")

    section.add_to(doc)

    doc.save("/Users/reinder.vosdewael/test.docx")

    assert doc.paragraphs[0].runs[0].text == "Hello world!"
    assert doc.paragraphs[0].style.name == "Heading 1"  # type: ignore[union-attr]


def test_paragraph_section_add_to_subsections() -> None:
    """Test adding a ParagraphSection with a subsection to a document."""
    doc = docx.Document()
    section = reports.ParagraphSection(
        content="Hello world!",
        style=cmi_docx.ParagraphStyle(bold=True),
        subsections=[reports.ParagraphSection(content="Goodbye world!", style=None)],
    )

    section.add_to(doc)

    assert doc.paragraphs[0].runs[0].text == "Hello world!"
    assert doc.paragraphs[0].runs[0].bold
    assert doc.paragraphs[1].runs[0].text == "Goodbye world!"
    assert doc.paragraphs[1].runs[0].style.name == "Default Paragraph Font"


def test_paragraph_section_with_condition_false() -> None:
    """Test adding a ParagraphSection with condition=False to a document."""
    doc = docx.Document()
    section = reports.ParagraphSection(
        content="Hello world!", style="Heading 1", condition=lambda: False
    )

    section.add_to(doc)

    assert len(doc.paragraphs) == 0


def test_table_section_basic(mock_sql_calls: None) -> None:
    """Test adding a basic TableSection to a document."""
    # The default Pyrite table style isn't present in python-docx' default
    # document. Use our template instead.
    pyrite_template = settings.DATA_DIR / "pyrite_template.docx"
    doc = docx.Document(str(pyrite_template))
    section = reports.TableSection(
        title="Test Table", level=2, tables=[scq.ScqTable("")]
    )

    section.add_to(doc)

    assert doc.paragraphs[-2].text == "Test Table"
    assert doc.paragraphs[-2].style.name == "Heading 2"  # type: ignore[union-attr]
    assert len(doc.tables) == 1
