"""Tests for the Pyrite tables base module."""

import functools

import cmi_docx
import docx
import pytest
from docx import document, table

from ctk_functions.routers.pyrite import types
from ctk_functions.routers.pyrite.tables import base

BOLD_TABLE_STYLE = cmi_docx.CellStyle(cmi_docx.ParagraphStyle(bold=True))


@pytest.fixture
def doc() -> document.Document:
    """Fixture for a docx document."""
    return docx.Document()


@pytest.fixture
def tbl(doc: document.Document) -> table.Table:
    """Fixture for a docx table."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "0,0"
    tbl.rows[1].cells[0].text = "1,0"
    tbl.rows[0].cells[1].text = "0,1"
    tbl.rows[1].cells[1].text = "1,1"
    return tbl  # type: ignore[no-any-return]


def test_conditional_cell_style(tbl: table.Table) -> None:
    """Tests the conditional cell style happy path."""
    style = base.ConditionalCellStyle(
        condition=lambda text: text == "1,1",
        style=BOLD_TABLE_STYLE,
    )

    style.apply([*tbl.rows[0].cells, *tbl.rows[1].cells])

    assert not tbl.rows[0].cells[0].paragraphs[0].runs[0].bold
    assert not tbl.rows[1].cells[0].paragraphs[0].runs[0].bold
    assert not tbl.rows[0].cells[1].paragraphs[0].runs[0].bold
    assert tbl.rows[1].cells[1].paragraphs[0].runs[0].bold


def test_conditional_table_style(tbl: table.Table) -> None:
    """Tests the conditional table style happy path."""
    style = base.ConditionalTableStyle(
        condition=lambda _, row, __: row == 0,
        style=BOLD_TABLE_STYLE,
    )

    style.apply(tbl, 0, 0)
    style.apply(tbl, 1, 0)

    assert tbl.rows[0].cells[0].paragraphs[0].runs[0].bold
    assert not tbl.rows[1].cells[0].paragraphs[0].runs[0].bold


@pytest.mark.parametrize(
    ("value", "low", "high", "low_inclusive", "high_inclusive", "expected"),
    [
        # Test two-sided
        (1, 0, 2, False, False, True),
        (-1, 0, 2, False, False, False),
        (3, 0, 2, False, False, False),
        # Test string input
        ("1", 0, 2, False, False, True),
        # Test one-sided
        (1, 0, None, False, False, True),
        (1, 2, None, False, False, False),
        (1, None, 2, False, False, True),
        (1, None, 0, False, False, False),
        # Test equality in low/high inclusive
        (0, 0, None, False, False, False),
        (0, 0, None, True, False, True),
        (0, None, 0, False, False, False),
        (0, None, 0, False, True, True),
    ],
)
def test_clinical_relevance_in_range(  # noqa: PLR0913
    value: float,
    low: float | None,
    high: float | None,
    *,
    low_inclusive: bool,
    high_inclusive: bool,
    expected: bool,
) -> None:
    """Tests the check whether a value is in the relevance range."""
    relevance = base.ClinicalRelevance(
        low=low,
        high=high,
        label=None,
        low_inclusive=low_inclusive,
        high_inclusive=high_inclusive,
        style=BOLD_TABLE_STYLE,
    )

    actual = relevance.in_range(value)

    assert actual == expected


@pytest.mark.parametrize(
    ("low", "high", "label", "low_inclusive", "high_inclusive", "expected"),
    [
        (0, 2, None, False, False, ">0, <2"),
        (0, 2, None, True, True, ">=0, <=2"),
        (0, 2, "Terra", False, False, ">0, <2 = Terra"),
        (0.010, 2.350, None, False, False, ">0.01, <2.35"),
        (0, None, None, False, False, ">0"),
        (None, 0, None, False, False, "<0"),
    ],
)
def test_clinical_relevance_string(  # noqa: PLR0913
    low: float | None,
    high: float | None,
    label: str | None,
    *,
    low_inclusive: bool,
    high_inclusive: bool,
    expected: str,
) -> None:
    """Tests conversion of clinical relevance to a string."""
    relevance = base.ClinicalRelevance(
        low=low,
        high=high,
        label=label,
        low_inclusive=low_inclusive,
        high_inclusive=high_inclusive,
        style=BOLD_TABLE_STYLE,
    )

    actual = str(relevance)

    assert actual == expected


def test_clinical_relevance_no_low_or_high() -> None:
    """Tests that an error is thrown for an invalid clinical relevance."""
    with pytest.raises(ValueError, match="Low or high must be defined."):
        base.ClinicalRelevance(
            low=None,
            high=None,
            label=None,
            style=BOLD_TABLE_STYLE,
        )


def test_paragraph_block(doc: document.Document) -> None:
    """Tests adding a new paragraph block."""
    block = base.ParagraphBlock(
        content="Hello World!",
        style=cmi_docx.ParagraphStyle(font_rgb=(255, 0, 0)),
        level=1,
    )

    para = block.add_to(doc)

    assert para.text == "Hello World!"
    assert para.style.name == "Heading 1"  # type: ignore[union-attr]
    assert para.runs[0].font.color.rgb == (255, 0, 0)


def test_formatter_styles(tbl: table.Table) -> None:
    """Tests whether the formatter correctly calls conditional styles."""
    style = base.ConditionalCellStyle(
        condition=lambda _text: True, style=BOLD_TABLE_STYLE
    )
    formatter = base.Formatter(conditional_cell_styles=[style])

    formatter.format(tbl, 0, 0)

    assert tbl.rows[0].cells[0].paragraphs[0].runs[0].bold


def test_word_table_markup_not_2d_array() -> None:
    """Tests whether a word table markup raises when not 2D array."""
    with pytest.raises(ValueError, match="All rows must have the same length."):
        base.WordTableMarkup(rows=[[base.WordTableCell(content="a")], []])


def test_word_table_section_add_to_mixin_faulty_protocol() -> None:
    """Tests whether a WordTableSectionAddToMixin with a faulty class errors."""

    class ShouldError(base.WordTableSectionAddToMixin):
        pass

    error_message = (
        "Classes using the AddToMixin must be a valid implementation of "
        "the AddToProtocol"
    )

    with pytest.raises(TypeError, match=error_message):
        ShouldError().add_to(None)  # type: ignore[arg-type]


def test_add_to_procotol() -> None:
    """Tests whether AddToProtocol isinstance() works."""

    class NotValid:
        pass

    class DataSource(base.DataProducer):
        @classmethod
        @functools.lru_cache
        def fetch(cls, mrn: str) -> tuple[tuple[str, ...], ...]:  # noqa: ARG003
            return ((),)

        @classmethod
        def test_ids(cls, mrn: str) -> tuple[types.TestId, ...]:  # noqa: ARG003
            return ()

    class Valid:
        mrn = "a"
        data_source = DataSource()
        formatters = ((base.Formatter(),),)

    assert not isinstance(NotValid, base._AddToProtocol)
    assert isinstance(Valid, base._AddToProtocol)
