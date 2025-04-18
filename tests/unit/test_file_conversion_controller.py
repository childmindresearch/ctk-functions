"""Tests for the file conversion controller."""

import pathlib

import cmi_docx
import docx
from docx.text import run

from ctk_functions.routers.file_conversion import controller


def is_run_font_color_red(run: run.Run) -> bool:
    """Checks if the font color of a run is red."""
    return '<w:color w:val="FF0000"' in str(run.font.color.element.xml)


def test_mark_warnings_as_red(tmp_path: pathlib.Path) -> None:
    """Tests marking warning labels as red."""
    filename = tmp_path / "test.docx"
    doc = docx.Document()
    sentences = [
        "This is a {{!WARNING-TEXT}}.",
        "Another {{!WARNING-TEXT-2}} {{!WARNING-TEXT}} here.",
    ]
    doc.add_paragraph(sentences[0])
    doc.add_paragraph(sentences[1])
    doc.save(str(filename))

    controller._mark_warnings_as_red(doc)

    assert doc.paragraphs[0].text == sentences[0]
    assert doc.paragraphs[1].text == sentences[1]
    assert doc.paragraphs[0].runs[1].text == "{{!WARNING-TEXT}}"
    assert doc.paragraphs[1].runs[1].text == "{{!WARNING-TEXT-2}}"
    assert doc.paragraphs[1].runs[3].text == "{{!WARNING-TEXT}}"
    assert not is_run_font_color_red(doc.paragraphs[0].runs[0])
    assert is_run_font_color_red(doc.paragraphs[0].runs[1])
    assert is_run_font_color_red(doc.paragraphs[1].runs[1])
    assert is_run_font_color_red(doc.paragraphs[1].runs[3])


def test_markdown2docx(tmp_path: pathlib.Path) -> None:
    """Tests the conversion of Markdown to docx."""
    markdown = "\n\n".join(  # noqa: FLY002
        [
            "# Header",
            "This is a paragraph.",
            "This is a {{!WARNING-TEXT}}.",
        ],
    )

    docx_bytes = controller.markdown2docx(
        markdown,
        formatting=cmi_docx.ParagraphStyle(bold=True),
    )
    filename = tmp_path / "test.docx"
    with filename.open("wb") as file:
        file.write(docx_bytes)
    doc = docx.Document(str(filename))

    assert doc.paragraphs[0].text == "Header"
    assert doc.paragraphs[1].text == "This is a paragraph."
    assert doc.paragraphs[2].runs[1].text == "WARNING-TEXT"
    assert not is_run_font_color_red(doc.paragraphs[2].runs[0])
    assert is_run_font_color_red(doc.paragraphs[2].runs[1])
    assert doc.paragraphs[0].runs[0].font.bold


def test_markdown2docx_lua(tmp_path: pathlib.Path) -> None:
    """Tests the custom lua filters."""
    markdown = "++underlined++\n\n|ttabbed"

    docx_bytes = controller.markdown2docx(markdown)
    filename = tmp_path / "test.docx"
    with filename.open("wb") as file:
        file.write(docx_bytes)
    doc = docx.Document(str(filename))

    assert doc.paragraphs[0].runs[0].underline
    assert doc.paragraphs[1].text[0] == "\t"


def test_remove_warning_brackets() -> None:
    """Tests that warning brackets are correctly removed."""
    doc = docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.text = "Hello {{!REMOVE-THESE}} there. {{REMOVE-THESE}}"

    controller._remove_curly_brackets(doc)

    assert paragraph.text == "Hello REMOVE-THESE there. REMOVE-THESE"
