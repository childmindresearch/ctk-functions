"""Unit tests for the writer module."""

import dataclasses
import datetime
from typing import Literal

import docx
import pytest
import pytest_mock

from ctk_functions.routers.intake.intake_processing import parser_models, writer


@dataclasses.dataclass
class Language:
    """Basic replacement for parser.Language."""

    name: str
    fluency: Literal["fluent", "proficient", "conversational", "basic"]


@dataclasses.dataclass
class MockGuardian:
    """Basic replacement for parser.GuardianInformation."""

    title_name: str = "Mr. Lukas Fink"


@dataclasses.dataclass
class MockPatient:
    """Basic replacement for parser.PatientInformation."""

    full_name: str = "Lea Avatar"
    first_name: str = "Lea"
    date_of_birth: datetime.datetime = datetime.datetime(
        2015,
        1,
        1,
        tzinfo=datetime.UTC,
    )
    guardian: MockGuardian = dataclasses.field(default_factory=MockGuardian)
    age_gender_label: str = "girl"
    pronouns: list[str] = dataclasses.field(
        default_factory=lambda: ["she", "her", "her", "hers", "herself"],
    )


@dataclasses.dataclass
class MockIntake:
    """Basic replacement for parser.IntakeInformation."""

    patient: MockPatient = dataclasses.field(default_factory=MockPatient)
    date_of_intake: str = "03/03/2023"


def test_replace_patient_information() -> None:
    """Test that the method returns correctly formatted text."""
    intake = MockIntake()
    document = docx.Document()
    paragraph = document.add_paragraph("{{FULL_NAME}} is a {{PRONOUN_0}}.")
    paragraph.add_run(" {{PRONOUN_2}}")
    report_writer = writer.ReportWriter(intake, "gpt-4o")  # type: ignore[arg-type]
    report_writer.report.document = document
    expected = "Lea Avatar is a she. her"

    report_writer.replace_patient_information()
    actual = report_writer.report.document.paragraphs[0].text

    assert actual == expected


@pytest.mark.parametrize(
    ("languages", "expected"),
    [
        ([], ""),
        ([Language("English", "fluent")], "is fluent in English"),
        (
            [Language("English", "fluent"), Language("Spanish", "proficient")],
            "is fluent in English and proficient in Spanish",
        ),
        (
            [
                Language("English", "fluent"),
                Language("Spanish", "proficient"),
                Language("French", "conversational"),
            ],
            "is fluent in English, proficient in Spanish, and conversational in French",
        ),
        (
            [
                Language("English", "fluent"),
                Language("Spanish", "proficient"),
                Language("French", "conversational"),
                Language("German", "basic"),
            ],
            (
                "is fluent in English, proficient in Spanish, conversational in "
                "French, and has basic skills in German"
            ),
        ),
        (
            [
                Language("English", "fluent"),
                Language("Spanish", "fluent"),
                Language("French", "conversational"),
            ],
            "is fluent in English and Spanish and conversational in French",
        ),
        (
            [
                Language("English", "fluent"),
                Language("Spanish", "fluent"),
                Language("French", "fluent"),
            ],
            "is fluent in English, Spanish, and French",
        ),
    ],
)
def test__join_patient_languages(languages: Language, expected: str) -> None:
    """Test that the _join_patient_languages method returns correctly formatted text."""
    actual = writer.ReportWriter._join_patient_languages(languages)  # type: ignore[arg-type] # Tested in test_valid_language_replacement.

    assert actual == expected


def test_set_superscript() -> None:
    """Test that the set_superscript method returns correctly formatted text."""
    intake = MockIntake()
    report = writer.ReportWriter(intake, "gpt-4o")  # type: ignore[arg-type]
    document = docx.Document()
    document.add_paragraph("For the 1st time.")
    report.report.document = document

    report.set_superscripts()

    assert document.paragraphs[0].text == "For the 1st time."
    assert document.paragraphs[0].runs[0].text == "For the 1"
    assert document.paragraphs[0].runs[1].text == "st"
    assert document.paragraphs[0].runs[2].text == " time."
    assert document.paragraphs[0].runs[1].font.superscript


def test__family_psychiatric_history_get_endorsed_diagnoses_no_diagnoses() -> None:
    """Test that the method returns nothing for no diagnoses."""
    history = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]

    endorsed_diagnoses = history._get_endorsed_diagnoses([])

    assert endorsed_diagnoses == []


def test__family_psychiatric_history_get_endorsed_diagnoses_specified(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test that the method calls the LLM when history is specified."""
    diagnoses = parser_models.FamilyPsychiatricHistory(
        name="autism",
        first_degree=True,
        second_degree=True,
        third_degree=True,
        is_diagnosed=True,
        details="paternal grandfather",
    )

    history = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]
    history.llm = mocker.Mock()
    history.llm.classify_family_relatedness = mocker.Mock(return_value="")

    history._get_endorsed_diagnoses([diagnoses])

    history.llm.classify_family_relatedness.assert_called_once_with(diagnoses.details)


def test__family_psychiatric_history_get_endorsed_diagnoses_unspecified() -> None:
    """Tests that unspecified family psychiatric diagnoses returns correct text."""
    diagnoses = parser_models.FamilyPsychiatricHistory(
        name="autism",
        first_degree=True,
        second_degree=True,
        third_degree=True,
        is_diagnosed=True,
        details="",
    )
    history = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]

    endorsed_diagnoses = history._get_endorsed_diagnoses([diagnoses])

    assert endorsed_diagnoses == [
        "autism (1st degree relative, 2nd degree relative, and 3rd degree relative)"
    ]


def test__family_psychiatric_history_merge_diagnoses_always_merge() -> None:
    """Test that diagnoses are always merged when always_merge=True."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name="alcohol abuse",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=True,
            details="father",
        ),
        parser_models.FamilyPsychiatricHistory(
            name="substance abuse",
            first_degree=False,
            second_degree=True,
            third_degree=False,
            is_diagnosed=False,
            details="uncle",
        ),
    ]

    fph = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]
    merged = fph._merge_diagnoses(history)

    # Should have one merged diagnosis
    substance_abuse_diagnoses = [d for d in merged if d.name == "substance abuse"]
    assert len(substance_abuse_diagnoses) == 1

    merged_diagnosis = substance_abuse_diagnoses[0]
    assert merged_diagnosis.first_degree
    assert merged_diagnosis.second_degree
    assert not merged_diagnosis.third_degree
    assert merged_diagnosis.is_diagnosed
    assert merged_diagnosis.details == "father uncle"


def test__family_psychiatric_history_conditional_merge_same_diagnosis() -> None:
    """Test conditional merge when all diagnoses have same is_diagnosed value."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name="conduct disorder",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=True,
            details="brother",
        ),
        parser_models.FamilyPsychiatricHistory(
            name="oppositional defiant disorder",
            first_degree=False,
            second_degree=True,
            third_degree=False,
            is_diagnosed=True,
            details="cousin",
        ),
    ]

    fph = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]
    merged = fph._merge_diagnoses(history)

    # Should merge since both are diagnosed (is_diagnosed=True)
    merged_diagnoses = [
        d for d in merged if d.name == "oppositional defiant or conduct disorders"
    ]
    assert len(merged_diagnoses) == 1

    merged_diagnosis = merged_diagnoses[0]
    assert merged_diagnosis.first_degree is True
    assert merged_diagnosis.second_degree is True
    assert merged_diagnosis.is_diagnosed is True


def test__family_psychiatric_history_conditional_merge_different_diagnosis() -> None:
    """Test conditional merge when diagnoses have different is_diagnosed values."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name="conduct disorder",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=True,
            details="brother",
        ),
        parser_models.FamilyPsychiatricHistory(
            name="oppositional defiant disorder",
            first_degree=False,
            second_degree=True,
            third_degree=False,
            is_diagnosed=False,
            details="cousin",
        ),
    ]

    fph = writer._FamilyPsychiatricHistory(patient=None, llm=None)  # type: ignore[arg-type]
    merged = fph._merge_diagnoses(history)

    conduct_diagnoses = [d for d in merged if d.name == "conduct disorder"]
    odd_diagnoses = [d for d in merged if d.name == "oppositional defiant disorder"]
    assert len(conduct_diagnoses) == 1
    assert len(odd_diagnoses) == 1


def test__family_psychiatric_history_remove_diagnoses_keeps_diagnosed() -> None:
    """Test that diagnosed conditions are kept."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name="autism",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=True,
            details="",
        ),
    ]

    result = writer._FamilyPsychiatricHistory._remove_diagnoses(history)

    assert len(result) == 1
    assert result[0].name == "autism"


def test__family_psychiatric_history_remove_diagnoses_removes_undiagnosed() -> None:
    """Test that undiagnosed conditions not in ALWAYS_SHOW are removed."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name=next(iter(writer._DisorderConfig.POSITIVE_ONLY)),
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=False,
            details="",
        ),
    ]

    result = writer._FamilyPsychiatricHistory._remove_diagnoses(history)

    assert len(result) == 0


def test__family_psychiatric_history_remove_diagnoses_keeps_always_show() -> None:
    """Test that undiagnosed conditions in ALWAYS_SHOW are kept."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name=next(iter(writer._DisorderConfig.ALWAYS_SHOW)),
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=False,
            details="",
        ),
    ]

    result = writer._FamilyPsychiatricHistory._remove_diagnoses(history)

    assert len(result) == 1


def test__family_psychiatric_history_remove_diagnoses_unknown_disorder() -> None:
    """Test that unknown disorders raise KeyError."""
    history = [
        parser_models.FamilyPsychiatricHistory(
            name="unknown_disorder",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=True,
            details="",
        ),
    ]

    with pytest.raises(KeyError, match="Unknown disorder: 'unknown_disorder'"):
        writer._FamilyPsychiatricHistory._remove_diagnoses(history)


def test__family_psychiatric_history_write_history_unknown(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test writing when family history is unknown."""
    patient = mocker.Mock()
    patient.possessive_first_name = "John's"

    fph = writer._FamilyPsychiatricHistory(patient=patient, llm=None)  # type: ignore[arg-type]
    result = fph._write_history_unknown()

    assert result == "John's family psychiatric history was unknown."


def test__family_psychiatric_history_write_psychiatric_history_both_unknown(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test when both maternal and paternal history are unknown."""
    patient = mocker.Mock()
    patient.psychiatric_history.family_psychiatric_history.is_father_history_known = (
        False
    )
    patient.psychiatric_history.family_psychiatric_history.is_mother_history_known = (
        False
    )

    fph = writer._FamilyPsychiatricHistory(patient=patient, llm=None)  # type: ignore[arg-type]
    fph._write_history_unknown = mocker.Mock(return_value="History unknown")

    result = fph.write_psychiatric_history()
    assert result == "History unknown"


def test__family_psychiatric_history_write_psychiatric_history_no_endorsed(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test psychiatric history with no endorsed diagnoses."""
    # Setup patient mock
    patient = mocker.Mock()
    patient.first_name = "John"
    patient.possessive_first_name = "John's"
    patient.guardian.title_name = "Ms. Smith"
    patient.psychiatric_history.family_psychiatric_history.is_father_history_known = (
        True
    )
    patient.psychiatric_history.family_psychiatric_history.is_mother_history_known = (
        True
    )
    patient.psychiatric_history.family_psychiatric_history.diagnoses = [
        parser_models.FamilyPsychiatricHistory(
            name="depression",
            first_degree=True,
            second_degree=False,
            third_degree=False,
            is_diagnosed=False,
            details="",
        ),
    ]
    fph = writer._FamilyPsychiatricHistory(patient=patient, llm=None)  # type: ignore[arg-type]

    result = fph.write_psychiatric_history()

    assert result.startswith(
        "John's family history is largely unremarkable for psychiatric illnesses."
    )


def test__family_psychiatric_history_write_psychiatric_history_paternal_unknown(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test when paternal history is unknown."""
    # Setup patient mock
    patient = mocker.Mock()
    patient.first_name = "John"
    patient.possessive_first_name = "John's"
    patient.guardian.title_name = "Ms. Smith"
    patient.psychiatric_history.family_psychiatric_history.is_father_history_known = (
        False
    )
    patient.psychiatric_history.family_psychiatric_history.is_mother_history_known = (
        True
    )
    patient.psychiatric_history.family_psychiatric_history.diagnoses = []
    fph = writer._FamilyPsychiatricHistory(patient=patient, llm=None)  # type: ignore[arg-type]

    result = fph.write_psychiatric_history()

    assert (
        "Information regarding John's paternal family psychiatric history was unknown."
        in result
    )


def test__family_psychiatric_history_write_psychiatric_history_maternal_unknown(
    mocker: pytest_mock.MockerFixture,
) -> None:
    """Test when maternal history is unknown."""
    # Setup patient mock
    patient = mocker.Mock()
    patient.first_name = "John"
    patient.possessive_first_name = "John's"
    patient.guardian.title_name = "Ms. Smith"
    patient.psychiatric_history.family_psychiatric_history.is_father_history_known = (
        True
    )
    patient.psychiatric_history.family_psychiatric_history.is_mother_history_known = (
        False
    )
    patient.psychiatric_history.family_psychiatric_history.diagnoses = []
    fph = writer._FamilyPsychiatricHistory(patient=patient, llm=None)  # type: ignore[arg-type]

    result = fph.write_psychiatric_history()

    assert (
        "Information regarding John's maternal family psychiatric history was unknown."
        in result
    )
