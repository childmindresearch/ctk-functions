"""Unit tests for table utilities."""

import cmi_docx
import docx
import pytest
from docx import table

from ctk_functions.routers.pyrite.tables import utils


@pytest.fixture
def mock_table() -> table.Table:
    """Create a mock table for testing."""
    doc = docx.Document()
    return doc.add_table(rows=2, cols=3)  # type: ignore[no-any-return]


def test_add_header(mock_table: table.Table) -> None:
    """Test adding headers to a table."""
    headers = ["Header1", "Header2", "Header3"]
    utils.add_header(mock_table, headers)

    for i, header in enumerate(headers):
        assert mock_table.rows[0].cells[i].text == header


def test_set_index_column_name_or_merge_different_names(
    mock_table: table.Table,
) -> None:
    """Test setting index column with different names."""
    mock_table.rows[0].cells[0].text = "Label1"

    utils.set_index_column_name_or_merge(mock_table, "Label2", 1, 0)

    assert mock_table.rows[0].cells[0].text == "Label1"
    assert mock_table.rows[1].cells[0].text == "Label2"


def test_set_index_column_name_or_merge_same_names(mock_table: table.Table) -> None:
    """Test setting index column with same names causing merge."""
    mock_table.rows[0].cells[0].text = "Label1"

    utils.set_index_column_name_or_merge(mock_table, "Label1", 1, 0)
    mock_table.rows[0].cells[0].text = "Label2"

    assert mock_table.rows[1].cells[0].text == "Label2"


@pytest.mark.parametrize(
    ("score", "expected_qualifier"),
    [
        (59, "extremely low"),
        (60, "very low"),
        (69, "very low"),
        (70, "low"),
        (79, "low"),
        (80, "low average"),
        (89, "low average"),
        (90, "average"),
        (109, "average"),
        (110, "high average"),
        (119, "high average"),
        (120, "high"),
        (129, "high"),
        (130, "very high"),
        (139, "very high"),
        (140, "extremely high"),
    ],
)
def test_standard_score_to_qualifier(score: int, expected_qualifier: str) -> None:
    """Test converting standard scores to qualifiers."""
    qualifier = utils.standard_score_to_qualifier(score)
    assert qualifier == expected_qualifier


@pytest.mark.parametrize(
    ("score", "expected_percentile"),
    [
        (67, 1),
        (69, 2),
        (72, 3),
        (74, 4),
        (76, 5),
        (77, 6),
        (78, 7),
        (79, 8),
        (80, 9),
        (81, 10),
        (82, 12),
        (90, 25),
        (100, 50),
        (110, 75),
        (120, 91),
        (125, 95),
        (130, 98),
        (133, 99),
        (150, 99),
    ],
)
def test_standard_score_to_percentile(score: int, expected_percentile: int) -> None:
    """Test converting standard scores to percentiles."""
    percentile = utils.standard_score_to_percentile(score)
    assert percentile == expected_percentile


def test_clinical_relevance_init_with_valid_parameters() -> None:
    """Test initialization with valid parameters."""
    low = 10
    high = 20
    label = "Normal"

    actual = utils.ClinicalRelevance(
        low=low,
        high=high,
        label=label,
        style=cmi_docx.TableStyle(),
    )

    assert actual.low == low
    assert actual.high == high
    assert actual.label == "Normal"


def test_clinical_relevance_init_with_missing_low_high() -> None:
    """Test initialization without low/high parameters."""
    with pytest.raises(
        ValueError,
        match="At least one of low or high must not be None.",
    ):
        utils.ClinicalRelevance(
            low=None,
            high=None,
            label="",
            style=cmi_docx.TableStyle(),
        )


def test_clinical_relevance_init_with_invalid_low_high() -> None:
    """Test initialization without low/high parameters."""
    with pytest.raises(ValueError, match="Low must be lower than high."):
        utils.ClinicalRelevance(low=1, high=0, label="", style=cmi_docx.TableStyle())


@pytest.mark.parametrize(
    ("low", "high", "label", "expected"),
    [
        (10, 20, "LABEL", "10-20 = LABEL"),
        (10, None, "LABEL", ">10 = LABEL"),
        (None, 10, "LABEL", "<10 = LABEL"),
        (10, None, None, ">10"),
    ],
)
def test_clinical_relevance_strings(
    low: int,
    high: int,
    label: str | None,
    expected: str,
) -> None:
    """Tests the string representation of ClinicalRelevance."""
    cr = utils.ClinicalRelevance(
        low=low,
        high=high,
        label=label,
        style=cmi_docx.TableStyle(),
    )

    assert str(cr) == expected


@pytest.mark.parametrize(
    ("low", "high", "value", "expected"),
    [
        (1, 10, 5, True),
        (1, 10, 11, False),
        (1, 10, -1, False),
        (None, 10, 9, True),
        (None, 10, 10, True),
        (None, 10, 11, False),
        (1, None, 1, False),
        (1, None, -1, False),
        (1, None, 2, True),
    ],
)
def test_clinical_relevance_in_range(
    low: int | None,
    high: int | None,
    value: int,
    expected: bool,  # noqa: FBT001
) -> None:
    """Tests that the in_range function evaluates correctly."""
    cr = utils.ClinicalRelevance(
        low=low,
        high=high,
        label="",
        style=cmi_docx.TableStyle(),
    )

    assert cr.in_range(value) == expected
