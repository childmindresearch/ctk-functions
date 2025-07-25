"""Tests for the referral endpoints."""

import pathlib

import docx
from fastapi import status, testclient


def test_referral_post(client: testclient.TestClient, tmp_path: pathlib.Path) -> None:
    """Tests the POST markdown2docx endpoint."""
    request = {
        "tables": [
            {
                "title": "Test Title",
                "table": {
                    "col1": ["1", "2", "3"],
                    "col2": ["4", "5", "6"],
                },
            }
        ]
    }

    response = client.post(
        "/referral",
        json=request,
    )

    test_file = tmp_path / "response.docx"
    with test_file.open("wb") as file:
        file.write(response.content)
    doc = docx.Document(str(test_file))

    assert response.status_code == status.HTTP_200_OK
    assert doc.paragraphs[1].text == "Test Title"
    assert len(doc.tables) == 1
