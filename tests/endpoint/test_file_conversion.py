"""Endpoint tests for file conversion."""

import pathlib

import docx
from fastapi import status, testclient


def test_markdown2docx(client: testclient.TestClient, tmp_path: pathlib.Path) -> None:
    """Tests the POST markdown2docx endpoint."""
    markdown = "# Heading \n\n Hello, World!"

    response = client.post("/markdown2docx", json={"markdown": markdown})
    with (tmp_path / "file.docx").open("wb") as file:
        file.write(response.read())
        document = docx.Document(str(tmp_path / "file.docx"))

    assert response.status_code == status.HTTP_200_OK
    assert document.paragraphs[0].text == "Heading"
    assert document.paragraphs[0].style.name == "Heading 1"  # type: ignore[union-attr]
    assert document.paragraphs[1].text == "Hello, World!"
    assert document.paragraphs[1].style.name == "First Paragraph"  # type: ignore[union-attr]
