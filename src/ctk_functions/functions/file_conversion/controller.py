"""Functions for converting files between different formats."""

import pathlib
import re
import tempfile
from typing import Any

import cmi_docx
import docx
import pypandoc
from docx import document, shared
from docx.oxml import ns


def markdown2docx(
    markdown: str,
    formatting: None | dict[str, Any] | cmi_docx.ParagraphStyle = None,
) -> bytes:
    """Converts a Markdown document to a .docx file.

    Uses a custom lua filter to allow underlining text between two '++'.

    Args:
        markdown: The Markdown document.
        formatting: Formatting options, must abide by cmi_docx.ParagraphStyle arguments
            if it is provided as a dictionary.

    Returns:
        The .docx file as bytes.
    """
    underline_filter = pathlib.Path(__file__).parent / "lua" / "underline.lua"
    tab_filter = pathlib.Path(__file__).parent / "lua" / "tab.lua"
    with tempfile.NamedTemporaryFile(suffix=".docx") as docx_file:
        pypandoc.convert_text(
            markdown,
            "docx",
            format="commonmark_x",
            outputfile=docx_file.name,
            filters=[str(underline_filter), str(tab_filter)],
        )

        docx_file.seek(0)
        document = docx.Document(docx_file.name)
        mark_warnings_as_red(document)
        set_list_indentations(document)
        remove_curly_brackets(document)

        if formatting is not None:
            if isinstance(formatting, dict):
                formatting = cmi_docx.ParagraphStyle(**formatting)
            for paragraph in document.paragraphs:
                extend_paragraph = cmi_docx.ExtendParagraph(paragraph)
                extend_paragraph.format(formatting)
        document.save(docx_file.name)
        return docx_file.read()


def mark_warnings_as_red(doc: document.Document) -> None:
    """Marks warning values as red.

    We use {{!WARNING-TEXT}} as a values for warnings that should be marked red.

    Args:
        doc: The document object.
    """
    extend_document = cmi_docx.ExtendDocument(doc)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    warning_regex = re.compile(r"{{!.*?}}")
    matches = warning_regex.finditer(text)
    unique_matches = {match.group() for match in matches}

    for match in unique_matches:
        extend_document.replace(match, match, cmi_docx.RunStyle(font_rgb=(255, 0, 0)))


def set_list_indentations(doc: document.Document) -> None:
    """Sets list indentations for the clinical reports.

    Args:
        doc: The document object.
    """
    for paragraph in doc.paragraphs:
        if not paragraph.style or paragraph.style.name != "Compact":
            continue

        paragraph_property = paragraph._p.pPr  # noqa: SLF001
        if paragraph_property is None:
            continue

        number_property = paragraph_property.find(ns.qn("w:numPr"))
        if number_property is None:
            continue

        indentation_level = number_property.find(ns.qn("w:ilvl"))
        if indentation_level is None:
            continue

        level = int(indentation_level.get(ns.qn("w:val"))) + 1
        paragraph.paragraph_format.left_indent = shared.Inches(0.25) * level
        paragraph.paragraph_format.first_line_indent = shared.Inches(-0.25)


def remove_curly_brackets(doc: document.Document) -> None:
    """Removes {{! and }} substrings.

    Args:
        doc: The document object.
    """
    for paragraph in doc.paragraphs:
        extended_paragraph = cmi_docx.ExtendParagraph(paragraph)
        extended_paragraph.replace("{{!", "")
        extended_paragraph.replace("{{", "")
        extended_paragraph.replace("}}", "")
