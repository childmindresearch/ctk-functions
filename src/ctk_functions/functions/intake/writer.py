"""Contains report writing functionality for intake information."""

import asyncio
import enum
import itertools

import cmi_docx
import docx
from docx import document as docx_document
from docx.enum import table as enum_table
from docx.enum import text as enum_text
from docx.text import paragraph as docx_paragraph

from ctk_functions import config
from ctk_functions.functions.intake import llm, parser, transformers
from ctk_functions.functions.intake.utils import (
    language_utils,
    string_utils,
)

settings = config.get_settings()
DATA_DIR = settings.DATA_DIR
AZURE_BLOB_CONNECTION_STRING = settings.AZURE_BLOB_CONNECTION_STRING
PLACEHOLDER = "______"


logger = config.get_logger()


class RGB(enum.Enum):
    """Represents an RGB color code for specific sections."""

    UNRELIABLE = (247, 150, 70)
    LLM = (0, 0, 255)
    TESTING = (155, 187, 89)


class StyleName(enum.Enum):
    """The styles for the report."""

    HEADING_1 = "Heading 1"
    HEADING_2 = "Heading 2"
    HEADING_3 = "Heading 3"
    TITLE = "Title"
    NORMAL = "Normal"


class ReportWriter:
    """Writes a report for intake information."""

    def __init__(self, intake: parser.IntakeInformation) -> None:
        """Initializes the report writer.

        Args:
            intake: The intake information.
        """
        logger.debug("Initializing the report writer.")
        self.intake = intake
        self.report: docx_document.Document = docx.Document(
            str(DATA_DIR / "report_template.docx")
        )
        self.insert_before = next(
            paragraph
            for paragraph in self.report.paragraphs
            if "MENTAL STATUS EXAMINATION AND TESTING BEHAVIORAL OBSERVATIONS"
            in paragraph.text
        )

        self.llm = llm.Llm(
            intake.patient.first_name,
            intake.patient.pronouns,
        )

        if not self.insert_before:
            msg = "Insertion point not found in the report template."
            raise ValueError(msg)

    async def transform(self) -> None:
        """Transforms the intake information to a report."""
        logger.debug("Transforming the intake information to a report.")

        self.write_reason_for_visit()
        self.write_developmental_history()
        self.write_academic_history()
        self.write_social_history()
        self.write_psychiatric_history()
        self.write_medical_history()
        self.write_current_psychiatric_functioning()
        self.add_page_break()
        self.add_footer()

        self.replace_patient_information()
        self.add_signatures()
        await self.make_llm_edits()
        await self.apply_corrections()

    def replace_patient_information(self) -> None:
        """Replaces the patient information in the report."""
        logger.debug("Replacing patient information in the report.")
        replacements = {
            "full_name": self.intake.patient.full_name,
            "preferred_name": self.intake.patient.first_name,
            "date_of_birth": self.intake.patient.date_of_birth.strftime("%m/%d/%Y"),
            "reporting_guardian": self.intake.patient.guardian.title_name,
            "aged_gender": self.intake.patient.age_gender_label,
            "pronoun_0": self.intake.patient.pronouns[0],
            "pronoun_1": self.intake.patient.pronouns[1],
            "pronoun_2": self.intake.patient.pronouns[2],
            "pronoun_4": self.intake.patient.pronouns[4],
            "placeholder": PLACEHOLDER,
        }

        extendedDocument = cmi_docx.ExtendDocument(self.report)
        for template, replacement in replacements.items():
            template_formatted = "{{" + template.upper() + "}}"
            extendedDocument.replace(template_formatted, replacement)

    def write_reason_for_visit(self) -> None:
        """Writes the reason for visit to the end of the report."""
        logger.debug("Writing the reason for visit to the report.")
        patient = self.intake.patient
        handedness = patient.handedness
        iep = patient.education.individualized_educational_program
        past_diagnoses = patient.psychiatric_history.past_diagnoses.transform(
            short=True,
        )
        classroom = patient.education.classroom_type
        age_determinant = "an" if patient.age in (8, 18) else "a"

        placeholder_id = self.llm.run_text_with_parent_input(
            text=(
                f"""
                    {patient.guardian.title_name} attended the present
                    evaluation due to concerns regarding {PLACEHOLDER}.
                    The family is hoping for {PLACEHOLDER}.
                    The family learned of the study through {PLACEHOLDER}."
            """
            ),
            parent_input=f"""
                Reason for visit: {patient.reason_for_visit}.
                Hopes: {patient.hopes}.
                Learned of study: {patient.learned_of_study}.
            """,
            additional_instruction="""
                The placeholders should be replaced with at most one sentence.
                Further details will be provided later in the report, so include
                only the most pertinent information here.
            """,
        )

        if patient.education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(
                int(patient.education.grade),
            )
        else:
            grade_superscript = ""

        text = f"""
                At the time of enrollment, {patient.first_name} was {age_determinant}
                {patient.age}-year-old, {handedness} {patient.age_gender_label}
                {past_diagnoses}. {patient.first_name} was placed in a {classroom}
                {patient.education.grade}{grade_superscript}
                grade classroom at {patient.education.school_name}.
                {patient.first_name} {iep}. {patient.first_name} and
                {patient.pronouns[2]} {patient.guardian.relationship}
                {placeholder_id}
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("REASON FOR VISIT", StyleName.HEADING_1)
        self._insert(text)

    def write_developmental_history(self) -> None:
        """Writes the developmental history to the end of the report."""
        logger.debug("Writing the developmental history to the report.")
        self._insert("DEVELOPMENTAL HISTORY", StyleName.HEADING_1)
        self.write_prenatal_history()
        self.write_developmental_milestones()
        self.write_early_education()

    def write_prenatal_history(self) -> None:
        """Writes the prenatal and birth history of the patient to the report."""
        logger.debug("Writing the prenatal history to the report.")
        patient = self.intake.patient
        development = patient.development
        pregnancy_symptoms = development.birth_complications
        delivery = development.delivery
        delivery_location = development.delivery_location
        adaptability = development.adaptability
        duration_of_pregnancy = development.duration_of_pregnancy

        text = f"""
            {patient.guardian.title_name} reported {pregnancy_symptoms}.
            {patient.first_name} was born at
            {duration_of_pregnancy} of gestation with {delivery} at
            {delivery_location}. {patient.first_name} had {adaptability}
            during infancy and was {development.soothing_difficulty.name} to
            soothe.
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Prenatal and Birth History", StyleName.HEADING_2)
        self._insert(text)

    def write_developmental_milestones(self) -> None:
        """Writes the developmental milestones to the report."""
        logger.debug("Writing the developmental milestones to the report.")
        patient = self.intake.patient
        started_walking = patient.development.started_walking
        started_talking = patient.development.started_talking
        daytime_dryness = patient.development.daytime_dryness
        nighttime_dryness = patient.development.nighttime_dryness

        texts = [
            f"""
            {patient.first_name}'s achievement of social, language, fine and
            gross motor developmental milestones were within normal limits, as
            reported by {patient.guardian.title_name}.
            """,
            f"""{patient.first_name}
            {started_walking} and {started_talking}.
            {patient.pronouns[0].capitalize()} {daytime_dryness} and
            {nighttime_dryness}.
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        self._insert("Developmental Milestones", StyleName.HEADING_2)
        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[0], texts[0], {"font_rgb": RGB.UNRELIABLE.value}
        )

    def write_early_education(self) -> None:
        """Writes the early education information to the report."""
        logger.debug("Writing the early education information to the report.")
        patient = self.intake.patient
        development = patient.development

        reporting_guardian = patient.guardian.title_name
        early_intervention = development.early_intervention_age
        cpse = development.cpse_age

        text = f"""
            {reporting_guardian} reported that
            {patient.first_name} {early_intervention} and {cpse}.
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Early Educational Interventions", StyleName.HEADING_2)
        self._insert(text)

    def write_academic_history(self) -> None:
        """Writes the academic history to the end of the report."""
        logger.debug("Writing the academic history to the report.")
        self._insert("ACADEMIC AND EDUCATIONAL HISTORY", StyleName.HEADING_1)
        self.write_previous_testing()
        self.write_academic_history_table()
        self.write_past_educational_history()
        self.write_current_education()

    def write_previous_testing(self) -> None:
        """Writes the previous testing information to the report."""
        logger.debug("Writing the previous testing information to the report.")
        patient = self.intake.patient

        text = f"""
        {patient.first_name} has no history of previous psychoeducational
        evaluations./{patient.first_name} was evaluated by {PLACEHOLDER} in 20XX.
        Documentation of the results of the evaluation(s) were unavailable at
        the time of writing this report/ Notable results include:
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Previous Testing", StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(font_rgb=RGB.UNRELIABLE.value)

    def write_academic_history_table(self) -> None:
        """Writes the academic history table to the report."""
        logger.debug("Writing the academic history table to the report.")
        paragraph = self._insert("Name, Date of Assessment")
        cmi_docx.ExtendParagraph(paragraph).format(
            bold=True,
            alignment=enum_text.WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

        table = self.report.add_table(7, 4)
        self.insert_before._p.addprevious(table._tbl)  # noqa: SLF001
        table.style = "Table Grid"
        header_row = table.rows[0].cells

        header_texts = [
            "Domain/Index/Subtest",
            "Standard Score",
            "Percentile Rank",
            "Descriptor",
        ]
        for i, header in enumerate(header_texts):
            header_row[i].text = header
            header_row[i].width = 10
            cmi_docx.ExtendCell(header_row[i]).format(
                bold=True,
                background_rgb=(217, 217, 217),
                alignment=enum_text.WD_ALIGN_PARAGRAPH.CENTER,
            )
        for row in table.rows:
            row.height = 1
            row.height_rule = enum_table.WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                cmi_docx.ExtendCell(cell).format(
                    line_spacing=1,
                    space_after=0,
                    space_before=0,
                )

    def write_past_educational_history(self) -> None:
        """Writes the past educational history to the report."""
        logger.debug("Writing the educational history to the report.")
        patient = self.intake.patient

        placeholder_id = self.llm.run_with_list_input(
            patient.education.past_schools,
            additional_instruction="""
                Try to keep the text as concise as possible. THE TEXT SHOULD BE A
                SINGLE PARAGRAPH!

                Include only information that may be pertinent to the child's
                mental health. For example, if the child had strong negative
                experiences at a school, that would be important to include,
                whereas a teacher remarking that the child was average would not
                be relevant unless that was a significant change from other
                schools.
            """,
        )

        self._insert("Educational History", StyleName.HEADING_2)
        self._insert(placeholder_id)

    def write_current_education(self) -> None:
        """Writes the current educational history to the report."""
        logger.debug("Writing the current educational history to the report.")
        patient = self.intake.patient
        education = patient.education
        if education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(education.grade)
        else:
            grade_superscript = ""

        if education.school_functioning:
            placeholder_id = self.llm.run_text_with_parent_input(
                text=f"{patient.guardian.title_name} reported that {PLACEHOLDER}.",
                parent_input=f"""
                    Details on school functioning: {education.school_functioning}
                """,
            )
        else:
            placeholder_id = ""

        texts = [
            f"""
                {patient.first_name} is currently in the
                {education.grade}{grade_superscript} grade at
                {education.school_name}.""",
            f"""
                {patient.first_name} does/does not receive special
                education services and maintains/does not have an IEP
                allowing accommodations for/including {PLACEHOLDER}.
            """,
            f"""
                {patient.first_name}'s current academic performance was
                described as "{education.performance}" by
                {patient.guardian.title_name}, as they receive mostly
                {education.grades}. {placeholder_id}
            """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[1], texts[1], {"font_rgb": RGB.UNRELIABLE.value}
        )

    def write_social_history(self) -> None:
        """Writes the social history to the end of the report."""
        logger.debug("Writing the social history to the report.")
        self._insert("SOCIAL HISTORY", StyleName.HEADING_1)
        self.write_home_and_adaptive_functioning()
        self.write_social_functioning()

    def write_home_and_adaptive_functioning(self) -> None:
        """Writes the home and adaptive functioning to the report."""
        logger.debug("Writing the home and adaptive functioning to the report.")
        patient = self.intake.patient
        household = patient.household
        language_fluencies = self._join_patient_languages(self.intake.patient.languages)

        text_home = f"""
            {patient.first_name} lives in {household.city},
            {household.state}, with {household.members}.
            The {patient.guardian.parent_or_guardian}s are
            {household.guardian_marital_status}.
            {string_utils.join_with_oxford_comma(household.languages)} {"are" if
            len(household.languages) > 1 else "is"} spoken at home.
            {patient.language_spoken_best} is reportedly
            {patient.first_name}'s preferred language.
            {patient.first_name} {language_fluencies}.
        """
        if not household.home_functioning:
            text_adaptive = f"""
                {patient.guardian.title_name} denied any concerns with
                {patient.pronouns[2]} functioning in the home setting.
            """
        else:
            text_adaptive = self.llm.run_text_with_parent_input(
                text=f"""
                    Per {patient.guardian.title_name},
                    {patient.first_name} has a history of {PLACEHOLDER}.
                """,
                parent_input=f"""
                    Details on home functioning: {household.home_functioning}
                """,
            )

        text_home = string_utils.remove_excess_whitespace(text_home)
        text_adaptive = string_utils.remove_excess_whitespace(text_adaptive)

        self._insert("Home and Adaptive Functioning", StyleName.HEADING_2)
        self._insert(text_home)
        self._insert(text_adaptive)

    def write_social_functioning(self) -> None:
        """Writes the social functioning to the report."""
        logger.debug("Writing the social functioning to the report.")
        patient = self.intake.patient
        social_functioning = patient.social_functioning

        hobbies_id = self.llm.run_text_with_parent_input(
            text=f"""
                {patient.guardian.title_name} reported that {patient.pronouns[0]} has
                {PLACEHOLDER} friends in {patient.pronouns[2]} peer group.
                {patient.first_name}'s hobbies include {PLACEHOLDER}.
                """,
            parent_input=f"""
             Hobbies: {social_functioning.hobbies}
             Number of friends: {social_functioning.n_friends}.
             """,
        )

        text = f"""
            {patient.guardian.title_name} was pleased to describe
            {patient.first_name} as a (insert adjective e.g., affectionate)
            {patient.age_gender_label}. {hobbies_id}
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Social Functioning", StyleName.HEADING_2)
        self._insert(text)

    def write_psychiatric_history(self) -> None:
        """Writes the psychiatric history to the end of the report."""
        logger.debug("Writing the psychiatric history to the report.")
        self._insert("PSYCHRIATIC HISTORY", StyleName.HEADING_1)
        self.write_past_psychiatric_diagnoses()
        self.write_past_psychiatric_hospitalizations()
        self.write_past_therapeutic_interventions()
        self.write_past_self_injurious_behaviors_and_suicidality()
        self.write_past_aggressive_behaviors_and_homicidality()
        self.exposure_to_violence_and_trauma()
        self.administration_for_childrens_services_involvement()
        self.write_family_psychiatric_history()

    def write_past_psychiatric_diagnoses(self) -> None:
        """Writes the past psychiatric diagnoses to the report."""
        logger.debug("Writing the past psychiatric diagnoses to the report.")
        patient = self.intake.patient

        past_diagnoses = patient.psychiatric_history.past_diagnoses.base
        if not past_diagnoses:
            text = f"""
                {patient.guardian.title_name} denied any history of past psychiatric
                diagnoses for {patient.first_name}.
            """
            text = string_utils.remove_excess_whitespace(text)
        else:
            instructions = f"""
                The start of the first sentence should be, verbatim, as
                follows: "{patient.guardian.title_name} reported that
                {patient.first_name} was diagnosed with the following psychiatric
                diagnoses: "
            """

            text = self.llm.run_with_list_input(
                items=past_diagnoses,
                additional_instruction=instructions,
            )

        self._insert("Past Psychiatric Diagnoses", StyleName.HEADING_2)
        self._insert(text)

    def write_past_psychiatric_hospitalizations(self) -> None:
        """Writes the past psychiatric hospitalizations to the report."""
        logger.debug("Writing the past psychiatric hospitalizations to the report.")
        patient = self.intake.patient
        text = f"""
            {patient.guardian.title_name} denied any history of past psychiatric
            hospitalizations for {patient.first_name}.
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Past Psychiatric Hospitalizations", StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(font_rgb=RGB.UNRELIABLE.value)

    def administration_for_childrens_services_involvement(self) -> None:
        """Writes the ACS involvement to the report."""
        logger.debug("Writing the ACS involvement to the report.")

        text = self._write_basic_psychiatric_history(
            "ACS involvement", self.intake.patient.psychiatric_history.children_services
        )
        self._insert(
            "Administration for Children's Services (ACS) Involvement",
            StyleName.HEADING_2,
        )
        self._insert(text)

    def write_past_aggressive_behaviors_and_homicidality(self) -> None:
        """Writes the past aggressive behaviors and homicidality to the report."""
        logger.debug(
            "Writing the past aggressive behaviors and homicidality to the report."
        )

        text = self._write_basic_psychiatric_history(
            "homicidality or severe physically aggressive behaviors towards others",
            self.intake.patient.psychiatric_history.aggresive_behaviors,
        )

        self._insert(
            "Past Severe Aggressive Behaviors and Homicidality",
            StyleName.HEADING_2,
        )
        self._insert(text)

    def write_family_psychiatric_history(self) -> None:
        """Writes the family psychiatric history to the report."""
        logger.debug("Writing the family psychiatric history to the report.")

        instructions = """
            You may group diagnoses together with the given label if they have identical
             responses:
                1. Specific learning disorder
                    - specific learning disorder with impairment in mathematics
                    - specific learning disorder with impairment in reading
                    - specific learning disorder with impairment in written expression

                2. Oppositional defiant or conduct disorders
                    - Conduct disorder
                    - Oppositional defiant disorder

                3. Substance abuse:
                    - Alcohol abuse
                    - Substance abuse

                4. Anxiety disorders
                    - Generalized anxiety disorder
                    - Separation anxiety
                    - Social anxiety

            The following diagnoses must be omitted if negative:
                - Enuresis/Encopresis
                - Excoriation
                - Gender dysphoria
                - Intellectual disability
                - Language disorder
                - Personality disorder
                - Reactive attachment disorder
                - Selective mutism
                - Agoraphobia
                - Specific phobias
                - Tics/Tourette’s
        """

        text = (
            self.intake.patient.psychiatric_history.family_psychiatric_history.replace(
                transformers.ReplacementTags.PREFERRED_NAME.value,
                self.intake.patient.first_name,
            )
        )
        placeholder_id = self.llm.run_edit(text, instructions)

        self._insert("Family Psychiatric History", StyleName.HEADING_2)
        self._insert(placeholder_id)

    def write_past_therapeutic_interventions(self) -> None:
        """Writes the past therapeutic history to the report."""
        logger.debug("Writing the past therapeutic history to the report.")
        patient = self.intake.patient
        guardian = patient.guardian
        interventions = patient.psychiatric_history.therapeutic_interventions

        if not interventions:
            text = f"""
                    {patient.guardian.title_name} denied any history of therapeutic
                    interventions.
                """
            text = string_utils.remove_excess_whitespace(text)
        else:
            text = self.llm.run_with_list_input(
                items=interventions,
                additional_instruction=f"""
                    An example structure for this paragraph follows:

                    From {PLACEHOLDER} to {PLACEHOLDER},
                    {patient.first_name} engaged in therapy with
                    {PLACEHOLDER} due to {PLACEHOLDER} at a
                    frequency of {PLACEHOLDER}. {guardian.title_name}
                    described the treatment as {PLACEHOLDER}.
                    Treatment was ended due to {PLACEHOLDER}.
                    """,
            )

        self._insert("Past Therapeutic Interventions", StyleName.HEADING_2)
        self._insert(text)

    def write_past_self_injurious_behaviors_and_suicidality(self) -> None:
        """Writes the past self-injurious behaviors and suicidality to the report."""
        logger.debug(
            "Writing the past self-injurious behaviors and suicidality to the report."
        )

        text = self._write_basic_psychiatric_history(
            "serious self-injurious harm or suicidal ideation",
            self.intake.patient.psychiatric_history.self_harm,
        )

        self._insert(
            "Past Self-Injurious Behaviors and Suicidality",
            StyleName.HEADING_2,
        )
        self._insert(text)

    def exposure_to_violence_and_trauma(self) -> None:
        """Writes the exposure to violence and trauma to the report."""
        logger.debug("Writing the exposure to violence and trauma to the report.")

        text = self._write_basic_psychiatric_history(
            "violence or trauma",
            self.intake.patient.psychiatric_history.violence_and_trauma,
        )

        self._insert("Exposure to Violence and Trauma", StyleName.HEADING_2)
        self._insert(text)

    def write_medical_history(self) -> None:
        """Writes the medical history to the end of the report."""
        logger.debug("Writing the medical history to the report.")
        patient = self.intake.patient
        primary_care = patient.primary_care

        texts = [
            f"""
            {patient.first_name}'s medical history is unremarkable for
            significant medical conditions. {patient.pronouns[0].capitalize()} is not
            currently taking any medications for chronic medical conditions.""",
            f"""
            {primary_care.glasses_hearing_device}.
            {primary_care.prior_diseases}.
            """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        self._insert("MEDICAL HISTORY", StyleName.HEADING_1)
        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[0], texts[0], {"font_rgb": RGB.UNRELIABLE.value}
        )

    def write_clinical_summary_and_impressions(self) -> None:
        """Writes the clinical summary and impressions to the report."""
        logger.debug("Writing the clinical summary and impressions to the report.")
        patient = self.intake.patient
        gender = patient.age_gender_label

        text = f"""
            {patient.first_name} is a
            sociable/resourceful/pleasant/hardworking/etc. {gender} who
            participated in the Healthy Brain Network research project through
            the Child Mind Institute in the interest of participating in
            research/due to parental concerns regarding {PLACEHOLDER}.
        """
        text = string_utils.remove_excess_whitespace(text)

        heading = self._insert("CLINICAL SUMMARY AND IMPRESSIONS", StyleName.HEADING_1)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(heading).format(font_rgb=RGB.TESTING.value)
        cmi_docx.ExtendParagraph(paragraph).format(font_rgb=RGB.TESTING.value)

    def write_current_psychiatric_functioning(self) -> None:
        """Writes the current psychiatric functioning to the report.

        Note: this section mixes color codings. Color decorators are applied
        to the called functions instead.
        """
        logger.debug("Writing the current psychiatric functioning to the report.")
        self._insert("CURRENT PSYCHIATRIC FUNCTIONING", StyleName.HEADING_1)
        self.write_current_psychiatric_medications_intake()
        self.write_current_psychiatric_medications_testing()
        self.write_denied_symptoms()

    def write_current_psychiatric_medications_intake(self) -> None:
        """Writes the current psychiatric medications to the report."""
        logger.debug("Writing the current psychiatric medications to the report.")
        patient = self.intake.patient
        text = f"""
            {patient.first_name} is currently prescribed a daily/twice daily
            oral course of {PLACEHOLDER} for {PLACEHOLDER}.
            {patient.pronouns[0].capitalize()} is being treated by Doctortype,
            DoctorName, monthly/weekly/biweekly. The medication has been
            ineffective/effective.
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Current Psychiatric Medications", StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(font_rgb=RGB.UNRELIABLE.value)

    def write_current_psychiatric_medications_testing(self) -> None:
        """Writes the current psychiatric medications to the report."""
        logger.debug("Writing the current psychiatric medications to the report.")
        patient = self.intake.patient
        texts = [
            f"""
        [Rule out presenting diagnoses, using headlines and KSADS/DSM criteria.
        Examples of headlines include: Temper Outbursts (ending should include
        “{patient.guardian.title_name} denied any consistent patterns of
        irritability for {patient.first_name}" if applicable), Inattention
        and Hyperactivity, Autism-Related Symptoms, Oppositional Defiant
        Behaviors, etc.].""",
            "Establish a baseline first for temper outbursts",
            f"""
           (Ex:
        Though {patient.first_name} is generally a {PLACEHOLDER} child,
        {patient.pronouns[0]} continues to have difficulties with temper
        tantrums…).
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        paragraph = self._insert(texts[0])
        paragraph.add_run(texts[1])
        paragraph.runs[-1].bold = True
        paragraph.add_run(texts[2])
        cmi_docx.ExtendParagraph(paragraph).format(
            font_rgb=RGB.TESTING.value, italics=True
        )

    def write_denied_symptoms(self) -> None:
        """Writes the denied symptoms to the report."""
        logger.debug("Writing the denied symptoms to the report.")
        patient = self.intake.patient
        text = f"""
        {patient.guardian.title_name} and {patient.first_name} denied any
        current significant symptoms related to mood, suicidality, psychosis,
        eating, oppositional or conduct behaviors, substance abuse, autism,
        tics, inattention/hyperactivity, enuresis/encopresis, trauma, sleep,
        panic, anxiety or obsessive-compulsive disorders.
        """
        text = string_utils.remove_excess_whitespace(text)

        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(font_rgb=RGB.TESTING.value)

    async def apply_corrections(self) -> None:
        """Applies various grammatical and styling corrections."""
        logger.debug("Applying corrections to the report.")
        document_corrector = language_utils.DocumentCorrections(self.report)
        await document_corrector.correct()

    def add_signatures(self) -> None:
        """Adds the signatures to the report.

        Michael Milham's signature is placed in a different location than the
        other signatures. As such, it needs some custom handling.
        """
        logger.debug("Adding signatures to the report.")
        signatures = (DATA_DIR / "signatures").glob("*.png")
        document = cmi_docx.ExtendDocument(self.report)

        for signature in signatures:
            signatory = signature.stem.replace("_", " ")
            paragraph_index = next(
                index
                for index in range(len(self.report.paragraphs))
                if self.report.paragraphs[index].text.lower().startswith(signatory)
            )

            if signatory != "michael p. milham":
                paragraph_index -= 1
            document.insert_image(paragraph_index, signature)

            if signatory != "michael p. milham":
                document._insert_empty_paragraph(paragraph_index)

    async def make_llm_edits(self) -> None:
        """Makes edits to the report using a large language model."""
        logger.debug("Making edits to the report using a large language model.")
        extendedDocument = cmi_docx.ExtendDocument(self.report)
        replacements = await asyncio.gather(
            *[placeholder.replacement for placeholder in self.llm.placeholders]
        )
        ids = [placeholder.id for placeholder in self.llm.placeholders]
        for id, replacement in zip(ids, replacements):
            extendedDocument.replace(
                id, replacement.strip(), {"font_rgb": RGB.LLM.value}
            )

    def add_footer(self) -> None:
        """Adds a footer to the report."""
        logger.debug("Adding a footer to the report.")
        footer = cmi_docx.ExtendParagraph(self.report.sections[0].footer.paragraphs[0])
        footer.paragraph.text = "Font Colors: Template, Testing, Large Language Model"
        footer.replace("Template", "Template", {"font_rgb": RGB.UNRELIABLE.value})
        footer.replace("Testing", "Testing", {"font_rgb": RGB.TESTING.value})
        footer.replace(
            "Large Language Model", "Large Language Model", {"font_rgb": RGB.LLM.value}
        )

    def add_page_break(self) -> None:
        """Adds a page break to the report."""
        paragraph = self._insert("")
        paragraph.add_run().add_break(enum_text.WD_BREAK.PAGE)

    def _insert(
        self,
        text: str,
        style: StyleName = StyleName.NORMAL,
    ) -> docx_paragraph.Paragraph:
        """Inserts text at the insertion point.

        Given the current structure of the report, text insertions only occur
        in one location. This point can always be modified by altering the
        `insert_before` attribute.

        Args:
            text: The text to insert.
            style: The style of the text.

        Returns:
            The new paragraph.
        """
        insertion_index = next(
            index
            for index in range(len(self.report.paragraphs))
            if self.report.paragraphs[index].text == self.insert_before.text
        )
        return cmi_docx.ExtendDocument(self.report).insert_paragraph_by_text(
            insertion_index, text, style.value
        )

    @staticmethod
    def _join_patient_languages(languages: list[parser.Language]) -> str:
        """Joins the patient's languages."""
        fluency_groups = itertools.groupby(
            languages,
            key=lambda language: language.fluency,
        )
        fluency_dict = {
            fluency: [language.name for language in language_group]
            for fluency, language_group in fluency_groups
        }

        language_descriptions = [
            f"{fluency} in {string_utils.join_with_oxford_comma(fluency_dict[fluency])}"
            for fluency in ["fluent", "proficient", "conversational"]
            if fluency in fluency_dict
        ]
        prepend_is = len(language_descriptions) > 0
        if "basic" in fluency_dict:
            language_descriptions.append(
                (
                    "has basic skills in "
                    f"{string_utils.join_with_oxford_comma(fluency_dict['basic'])}"
                ),
            )

        text = string_utils.join_with_oxford_comma(language_descriptions)
        if prepend_is:
            text = "is " + text
        return text

    def _write_basic_psychiatric_history(self, label: str, report: str) -> str:
        """Writes the ACS involvement to the report."""
        if not report:
            text = f"""
                {self.intake.patient.guardian.title_name} denied any history of
                {label} for {self.intake.patient.first_name}.
            """
            return string_utils.remove_excess_whitespace(text)
        return self.llm.run_text_with_parent_input(
            f"{self.intake.patient.guardian.title_name} reported that {PLACEHOLDER}.",
            parent_input=f"Details: {report}",
        )
