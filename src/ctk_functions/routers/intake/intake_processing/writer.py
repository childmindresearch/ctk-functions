"""Contains report writing functionality for intake information."""

import asyncio
import bisect
import copy
import dataclasses
import enum
import itertools
import re

import cmi_docx
import docx
import pydantic
from cmi_docx import comment, styles
from docx.enum import table as enum_table
from docx.enum import text as enum_text
from docx.text import paragraph as docx_paragraph

from ctk_functions.core import config, word
from ctk_functions.microservices import redcap
from ctk_functions.routers.intake.intake_processing import (
    parser,
    parser_models,
    writer_llm,
)
from ctk_functions.routers.intake.intake_processing.utils import (
    language_utils,
    string_utils,
)

settings = config.get_settings()
DATA_DIR = settings.DATA_DIR
PLACEHOLDER = "______"
COMMENT_AUTHOR = "Clinician Toolkit"


logger = config.get_logger()


class _RGB(enum.Enum):
    """Represents an RGB color code for specific sections."""

    BASIC = (0, 0, 0)
    LLM = (0, 0, 255)
    TESTING = (155, 187, 89)
    UNRELIABLE = (247, 150, 70)
    ERROR = (255, 0, 0)


class EnabledTasks(pydantic.BaseModel):
    """Enables or disables subsets of the intake processing.

    Used only in developer testing, these should all be True in production.
    """

    model_config = pydantic.ConfigDict(frozen=True)

    reason_for_visit: bool = True
    developmental_history: bool = True
    academic_history: bool = True
    social_history: bool = True
    psychiatric_history: bool = True
    medical_history: bool = True
    current_psychiatric_functioning: bool = True
    corrections: bool = True
    signatures: bool = True


class ReportWriter:
    """Writes a report for intake information."""

    def __init__(
        self,
        intake: parser.IntakeInformation,
        enabled_tasks: EnabledTasks | None = None,
    ) -> None:
        """Initializes the report writer.

        Args:
            intake: The intake information.
            enabled_tasks: Defines which subsets of the intake form to run.
        """
        logger.debug("Initializing the report writer.")
        self.intake = intake
        self.report = cmi_docx.ExtendDocument(
            docx.Document(str(DATA_DIR / "report_template.docx")),
        )
        self.enabled_tasks = enabled_tasks or EnabledTasks()
        self.insert_before = next(
            paragraph
            for paragraph in self.report.document.paragraphs
            if "MENTAL STATUS EXAMINATION AND TESTING BEHAVIORAL OBSERVATIONS"
            in paragraph.text
        )

        self.llm = writer_llm.WriterLlm(
            intake.patient.first_name,
            intake.patient.pronouns,
        )

        if not self.insert_before:
            msg = "Insertion point not found in the report template."
            raise ValueError(msg)

    async def transform(self) -> None:
        """Transforms the intake information to a report."""
        logger.debug("Transforming the intake information to a report.")

        if self.enabled_tasks.reason_for_visit:
            self.write_reason_for_visit()
        if self.enabled_tasks.developmental_history:
            self.write_developmental_history()
        if self.enabled_tasks.academic_history:
            self.write_academic_history()
        if self.enabled_tasks.social_history:
            self.write_social_history()
        if self.enabled_tasks.psychiatric_history:
            self.write_psychiatric_history()
        if self.enabled_tasks.medical_history:
            self.write_medical_history()
        if self.enabled_tasks.current_psychiatric_functioning:
            self.write_current_psychiatric_functioning()

        self.add_page_break()
        self.add_footer()
        self.replace_patient_information()

        if self.enabled_tasks.signatures:
            self.add_signatures()

        promises = [
            placeholder.await_replacement() for placeholder in self.llm.placeholders
        ]
        if self.enabled_tasks.corrections:
            promises.append(self.apply_corrections())
        # Python async execution is lazy, so we try to execute as many async calls
        # concurrently in a single asyncio.gather.
        await asyncio.gather(*promises)
        self.make_llm_edits()
        self.set_superscripts()

    def replace_patient_information(self) -> None:
        """Replaces the patient information in the report."""
        logger.debug("Replacing patient information in the report.")

        basic_style = styles.RunStyle(font_rgb=_RGB.BASIC.value)
        replacements = {
            "full_name": (self.intake.patient.full_name, basic_style),
            "date_of_intake": (
                self.intake.date_of_intake or "XX/XX/XXXX",
                (
                    basic_style
                    if self.intake.date_of_intake
                    else styles.RunStyle(font_rgb=_RGB.UNRELIABLE.value)
                ),
            ),
            "preferred_name": (self.intake.patient.first_name, basic_style),
            "date_of_birth": (
                self.intake.patient.date_of_birth.strftime("%m/%d/%Y"),
                basic_style,
            ),
            "reporting_guardian": (
                self.intake.patient.guardian.title_name,
                basic_style,
            ),
            "aged_gender": (self.intake.patient.age_gender_label, basic_style),
            "pronoun_0": (self.intake.patient.pronouns[0], basic_style),
            "pronoun_1": (self.intake.patient.pronouns[1], basic_style),
            "pronoun_2": (self.intake.patient.pronouns[2], basic_style),
            "pronoun_4": (self.intake.patient.pronouns[4], basic_style),
            "placeholder": (PLACEHOLDER, basic_style),
        }

        for template, replacement in replacements.items():
            template_formatted = "{{" + template.upper() + "}}"
            self.report.replace(
                template_formatted,
                replacement[0],
                style=replacement[1],
            )

    def write_reason_for_visit(self) -> None:
        """Writes the reason for visit to the end of the report."""
        logger.debug("Writing the reason for visit to the report.")
        patient = self.intake.patient
        past_diagnoses = patient.psychiatric_history.past_diagnoses.transform(
            short=True,
        )
        age_determinant = "an" if patient.age in (8, 11, 18) else "a"

        if patient.education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(
                int(patient.education.grade),
            )
        else:
            grade_superscript = ""

        main_text = string_utils.remove_excess_whitespace(f"""
            At the time of enrollment, {patient.first_name} was
            {age_determinant} {patient.age}-year-old, {patient.handedness}
            {patient.age_gender_label} {past_diagnoses}.
            {patient.first_name} was placed in a {patient.education.classroom_type}
            {patient.education.grade}{grade_superscript} grade classroom at
            {patient.education.school_name}. {patient.first_name}
            {patient.education.individualized_educational_program}.
            {patient.first_name} and {patient.pronouns[2]}
            {patient.guardian.relationship}
        """)

        hopes_text = (
            "The family is hoping for diagnostic clarity and "
            "recommendations on an appropriate course of interventions."
        )

        attendance_id = self.llm.run_text_with_parent_input(
            text=f"""
                {patient.guardian.title_full_name} attended the present
                evaluation due to concerns regarding {PLACEHOLDER}.
            """,
            parent_input=f"Reason for visit: {patient.reason_for_visit}.",
            additional_instruction="""
                The placeholder should be replaced with at most one sentence.
                Further details will be provided later in the report, so include
                only the most pertinent information here.
            """,
            context=main_text,
            comment=f"Reason for visit: {patient.reason_for_visit}",
        )

        learned_of_study_id = self.llm.run_text_with_parent_input(
            text=f"The family learned of the study through {PLACEHOLDER}.",
            parent_input=f"Learned of study: {patient.learned_of_study}.",
            additional_instruction="""
                The placeholders should be replaced with at most one sentence.
                Further details will be provided later in the report, so include
                only the most pertinent information here.
            """,
            context=main_text + " " + hopes_text,
            comment=f"Learned of study: {patient.learned_of_study}.",
        )

        self._insert("REASON FOR VISIT", word.StyleName.HEADING_1)
        paragraph = self._insert(main_text)
        paragraph.add_run(" " + attendance_id + " ")
        hopes_run = paragraph.add_run(hopes_text)

        comment.add_comment(
            self.report.document,
            location=(hopes_run, hopes_run),
            author=COMMENT_AUTHOR,
            text=f"Hopes: {patient.hopes}",
        )

        paragraph.add_run(" " + learned_of_study_id)
        self._insert("")

        # Due to a known error in the intake form multiple classifications may appear.
        # Mark these in red as this should not be possible.
        if len(patient.education.iep_classifications) > 1:
            cmi_docx.ExtendParagraph(paragraph).replace(
                "/".join(patient.education.iep_classifications),
                "/".join(patient.education.iep_classifications),
                cmi_docx.RunStyle(font_rgb=_RGB.ERROR.value),
            )

    def write_developmental_history(self) -> None:
        """Writes the developmental history to the end of the report."""
        logger.debug("Writing the developmental history to the report.")
        self._insert("DEVELOPMENTAL HISTORY", word.StyleName.HEADING_1)
        self.write_prenatal_history()
        self.write_developmental_milestones()
        self.write_early_education_interventions()

    def write_prenatal_history(self) -> None:
        """Writes the prenatal and birth history of the patient to the report."""
        logger.debug("Writing the prenatal history to the report.")
        patient = self.intake.patient
        development = patient.development
        mother_name = patient.mother.title_name if patient.mother else "[MOTHER NAME]"

        location_str = str(development.delivery_location).strip()
        if location_str:
            location_str = f" {location_str}"

        birth_sentence = (
            f"{patient.first_name} was born at {development.duration_of_pregnancy} of "
            f"gestation with {development.delivery}{location_str}."
        )

        if development.prenatal_issues:
            prenatal_sentence = self.llm.run_edit(
                text=f"""
                    {mother_name}'s prenatal period was remarkable for
                    {development.prenatal_issues} and
                    "{development.prenatal_issues_description}.".
                """,
                comment=f"""
                    Reported issues: {development.prenatal_issues}.\n\n
                    Further description: {development.prenatal_issues_description}.
                """,
                additional_instruction=f"""The output should contain only a single
                sentence. Your goal is to merge the quoted text in naturally.
                Do not alter the '{mother_name}'s prenatal period was remarkable for'
                snippet. Report only complications that occurred i.e. do not use
                phrases like "with no other complications".""",
            )
        else:
            prenatal_sentence = (
                f"{patient.guardian.title_name} reported an "
                "uncomplicated prenatal period."
            )

        if development.delivery.base == redcap.BirthDelivery.cesarean:
            # Contains a quote by the parent.
            birth_sentence = self.llm.run_edit(birth_sentence, comment=birth_sentence)

        text = f"""
            {prenatal_sentence} {birth_sentence} {patient.first_name} had
            {development.adaptability} during infancy and was
            {development.soothing_difficulty.name} to soothe.
        """

        text = string_utils.remove_excess_whitespace(text)
        if development.infant_difficulties.any():
            placeholder = self.llm.run_with_object_input(
                items=[development.infant_difficulties],
                additional_instruction="""
                    You will receive a JSON of difficulties the child had during
                    infancy. Write text on this that can be appended to the paragraph
                    provided. Include existing issues i.e. do not include issues that
                    were reported as not existing. Do not include a phrase stating
                    that no other issues were reported or similar.
                """,
                verify=True,
                context=text,
                comment=str(development.infant_difficulties),
            )
            text += f" {placeholder}"

        self._insert("Prenatal and Birth History", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_developmental_milestones(self) -> None:
        """Writes the developmental milestones to the report."""
        logger.debug("Writing the developmental milestones to the report.")
        patient = self.intake.patient
        started_walking = patient.development.started_walking
        started_talking = patient.development.started_talking
        daytime_dryness = patient.development.daytime_dryness
        nighttime_dryness = patient.development.nighttime_dryness

        texts = [
            f"""
            {patient.first_name}'s achievement of social, language, fine and
            gross motor developmental milestones were within normal limits, as
            reported by {patient.guardian.title_name}.
            """,
            f"""{patient.first_name}
            {started_walking} and {started_talking}.
            {patient.pronouns[0].capitalize()} {daytime_dryness} and
            {nighttime_dryness}.
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        self._insert("Developmental Milestones", word.StyleName.HEADING_2)
        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[0],
            texts[0],
            cmi_docx.RunStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        self._insert("")

    def write_early_education_interventions(self) -> None:
        """Writes the early education intervention information to the report."""
        logger.debug("Writing the early education information to the report.")
        patient = self.intake.patient
        development = patient.development

        reporting_guardian = patient.guardian.title_name
        early_intervention = development.early_intervention
        cpse = development.cpse_services

        if early_intervention + cpse:
            text = self.llm.run_with_object_input(
                items=early_intervention + cpse,
                additional_instruction=f"""
                    You will receive a list of all CPSE and early intervention services
                    received.

                    You should write a paragraph for the reporting of early
                    intervention and committee on preschool special education (CPSE)
                    services. An example paragraph structure follows:

                    "{reporting_guardian} reported that {patient.first_name}
                    received occupational therapy (2x/monthly) from May 2020 - June
                    2021  and speech therapy (1x/weekly) from June 2022 - July 2022.
                    services. {reporting_guardian} denied any history of Early
                    Intervention for {patient.first_name}."

                    Make sure to always report dates as Month Year - Month Year, unless
                    the service is ongoing, or the dates are unknown. Always put the
                    frequency between brackets.

                    If no EI was received, write that the reporting
                    guardian denied any history of early intervention.

                    If no CPSE services were received, write that the reporting guardian
                    denied any history of CPSE services.

                    If neither were received, write that the reporting guardian denied
                    any history of early intervention or CPSE services.

                    The first time that you write EI or CPSE, write it out in full
                    and put the acronym between brackets i.e. "Early Intervention (EI)".
                """,
                comment="\n\n".join(
                    [str(model) for model in early_intervention + cpse],
                ),
            )
        else:
            text = f"""
                {reporting_guardian} denied any history of Early Intervention (EI) or
                Committee on Preschool Special Education (CPSE) services for
                {patient.first_name}.
            """
            text = string_utils.remove_excess_whitespace(text)

        self._insert("Early Educational Interventions", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_academic_history(self) -> None:
        """Writes the academic history to the end of the report."""
        logger.debug("Writing the academic history to the report.")
        self._insert("ACADEMIC AND EDUCATIONAL HISTORY", word.StyleName.HEADING_1)
        self.write_previous_testing()
        self.write_academic_history_table()
        self.write_past_educational_history()
        self.write_current_education()

    def write_previous_testing(self) -> None:
        """Writes the previous testing information to the report."""
        logger.debug("Writing the previous testing information to the report.")
        patient = self.intake.patient

        text = f"""
        {patient.first_name} has no history of previous psychoeducational
        evaluations./{patient.first_name} was evaluated by {PLACEHOLDER} in 20XX.
        Documentation of the results of the evaluation(s) were unavailable at
        the time of writing this report/ Notable results include:
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Previous Testing", word.StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(
            cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        self._insert("")

    def write_academic_history_table(self) -> None:
        """Writes the academic history table to the report."""
        logger.debug("Writing the academic history table to the report.")
        paragraph = self._insert("Name, Date of Assessment")
        cmi_docx.ExtendParagraph(paragraph).format(
            cmi_docx.ParagraphStyle(
                bold=True,
                alignment=enum_text.WD_PARAGRAPH_ALIGNMENT.CENTER,
            ),
        )

        table = self.report.document.add_table(7, 4)
        self.insert_before._p.addprevious(table._tbl)  # noqa: SLF001
        table.style = "Table Grid"
        header_row = table.rows[0].cells

        header_texts = [
            "Domain/Index/Subtest",
            "Standard Score",
            "Percentile Rank",
            "Descriptor",
        ]
        for index, header in enumerate(header_texts):
            header_row[index].text = header
            header_row[index].width = 10
            cmi_docx.ExtendCell(header_row[index]).format(
                cmi_docx.CellStyle(
                    paragraph=cmi_docx.ParagraphStyle(
                        bold=True,
                        alignment=enum_text.WD_ALIGN_PARAGRAPH.CENTER,
                    ),
                    background_rgb=(217, 217, 217),
                ),
            )
        for row in table.rows:
            row.height = 1
            row.height_rule = enum_table.WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                cmi_docx.ExtendCell(cell).format(
                    cmi_docx.CellStyle(
                        paragraph=cmi_docx.ParagraphStyle(
                            line_spacing=1,
                            space_after=0,
                            space_before=0,
                        ),
                    ),
                )
        self._insert("")

    def write_past_educational_history(self) -> None:
        """Writes the past educational history to the report."""
        logger.debug("Writing the educational history to the report.")
        patient = self.intake.patient

        if patient.education.concerns is None:
            concern_instructions = ""
        else:
            concern_instructions = f"""
                {patient.guardian.title_name} reported the following overall concerns:
                "{patient.education.concerns}". Do your best to include these concerns
                where they chronologically make sense e.g. if concerns started in grade
                6, then include them with the school where the child was in grade 6.
                If no chronology can be derived, place these concerns at the start of
                the text.
            """

        placeholder_id = self.llm.run_with_object_input(
            patient.education.past_schools,
            additional_instruction=f"""
                {concern_instructions}

                Keep the text as concise as possible.

                Include only information that may be pertinent to the child's
                mental health. For example, if the child had strong negative
                experiences at a school, that would be important to include,
                whereas a teacher remarking that the child was average would not
                be relevant unless that was a significant change from other
                schools.
            """,
            comment="\n\n".join(
                [
                    f"{patient.guardian.title_name} reported the following overall "
                    f"concerns: {patient.education.concerns}",
                ]
                + [str(school) for school in patient.education.past_schools],
            ),
            verify=True,
        )

        self._insert("Educational History", word.StyleName.HEADING_2)
        self._insert(placeholder_id)
        self._insert("")

    def write_current_education(self) -> None:
        """Writes the current educational history to the report."""
        logger.debug("Writing the current educational history to the report.")
        patient = self.intake.patient
        education = patient.education
        if education.grade.isnumeric():
            grade_superscript = string_utils.ordinal_suffix(education.grade)
        else:
            grade_superscript = ""

        texts = [
            f"""
                At the time of intake, {patient.first_name} was in the
                {education.grade}{grade_superscript} grade at
                {education.school_name}.""",
        ]
        if education.iep_services:
            services_text = self.llm.run_with_object_input(
                education.iep_services,
                comment=", ".join(str(service) for service in education.iep_services),
                context=texts[0],
                additional_instruction=f"""
                    Write the IEP services that the child had access to.
                    If no date or frequency are reported, do not mention
                    that these are not reported; simply omit this information.

                    The phrase should start with: "{patient.pronouns[0].capitalize()}
                    maintained an IEP allowing for "
                """,
            )
        else:
            services_text = f"""
                {patient.first_name} does not receive special
                education services.
            """

        if education.testing_accommodations:
            accommodation_text = self.llm.run_edit(
                f"{patient.first_name} maintains an IEP allowing accommodations for "
                + string_utils.oxford_comma(education.testing_accommodations),
                comment="Testing accommodations: "
                + string_utils.oxford_comma(education.testing_accommodations),
                context=" ".join(texts),
                additional_instruction="Use only one sentence.",
            )
        else:
            accommodation_text = ""

        texts.append(services_text)
        texts.append(
            f"""
                {accommodation_text} {patient.first_name}'s current academic performance
                was described as "{education.performance}" by
                {patient.guardian.title_name},
                {education.grades}.
            """,
        )

        parent_reports = {
            "school functioning": education.school_functioning,
            "detention": education.detention,
            "suspensions": education.suspensions,
        }
        parent_reports = {key: value for key, value in parent_reports.items() if value}
        if parent_reports:
            placeholder_id = self.llm.run_text_with_parent_input(
                text=f"{patient.guardian.title_name} reported that {PLACEHOLDER}.",
                parent_input=str(parent_reports),
                context=" ".join(texts),
                comment=(
                    f"{patient.guardian.title_name} reported: "
                    + "\n\n".join(
                        f"{key}: {value!s}" for key, value in parent_reports.items()
                    )
                ),
            )
        else:
            placeholder_id = ""

        texts[-1] += f" {placeholder_id}"
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[1],
            texts[1],
            cmi_docx.RunStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        self._insert("")

    def write_social_history(self) -> None:
        """Writes the social history to the end of the report."""
        logger.debug("Writing the social history to the report.")
        self._insert("SOCIAL HISTORY", word.StyleName.HEADING_1)
        self.write_home_and_adaptive_functioning()
        self.write_social_functioning()

    def write_home_and_adaptive_functioning(self) -> None:
        """Writes the home and adaptive functioning to the report."""
        logger.debug("Writing the home and adaptive functioning to the report.")
        patient = self.intake.patient
        household = patient.household

        comment_text = (
            (
                "Household Members:\n{members}\n\nCity: {city}\nState: {state}"
                "\nHome Functioning:{functioning}\nHousehold Languages:{languages}\n\n"
            ).format(
                members="\n\n".join(str(member) for member in household.members),
                city=household.city,
                state=household.state,
                functioning=household.home_functioning,
                languages=", ".join([lang.name for lang in household.languages]),
            )
            + "Child Languages:\n"
            + "\n\n".join([str(lang) for lang in patient.languages])
        )
        text_home = self.llm.run_with_object_input(
            items={
                "household": household,
                "patients_languages": patient.languages,
            },
            additional_instruction=f"""
                You will receive a list of items, either descriptions of household
                members or the languages spoken by the patient, including details
                of where they speak them and their fluency.

                Your task is to write a paragraph for the home functioning of
                this patient.

                Languages spoken in the household are as follows:
                {
                string_utils.oxford_comma(
                    [language.name for language in household.languages]
                )
            }.

                What follows is an example output:

                "{patient.first_name} lives in {household.city}, {household.state},
                with {patient.pronouns[2]} biological parents,
                brother (age 10) and sister (age 15). The family is intact.
                {patient.first_name}'s mother, MOTHER FIRST AND LAST NAME (age), is a
                teacher, and {patient.pronouns[2]} father, FATHER FIRST
                AND LAST NAME (age), is a fireman. {patient.first_name} has
                mixed relationships with {patient.pronouns[2]} family members.

                The family maintains a bilingual household, speaking English and French.
                {patient.language_spoken_best} is reportedly
                {patient.first_name}'s preferred language.
                {patient.pronouns[2].capitalize()} level of proficiency in French is
                intermediate."

                If only one language is spoken at home, the final paragraph may be
                simplified to e.g. "English is the only language spoken in the home".

                Ensure you abide by the following rules:
                    - Do not include headers.
                    - Only the occupations of adults in the household should be included.
                    - Do not include grades for children.
                    - Summarize the relationship with the family as shortly as possible
                      even if they are different. For example  you may merge 'fair'
                      and 'excellent' relationships as 'positive'.
                    - Stick to the provided format wherein the children's ages are
                      mentioned in the first sentence, but adults' ages are mentioned
                      with their introduction.
                    - Ensure that the city ({household.city}) is included.
                    - Ensure that the state ({household.state}) is included.
                    - Do not include sibling names.
            """,  # noqa: E501
            verify=True,
            comment=comment_text,
        )

        if not household.home_functioning:
            text_adaptive = f"""
                {patient.guardian.title_name} denied any concerns regarding
                {patient.possessive_first_name} sensory processing, daily living skills
                or hygiene in the home.// OR Per {patient.guardian.title_name},
                {patient.first_name} has a history of difficulties with daily living
                skills and hygiene. ((include details of whether these tasks are not
                completed due to lack of willingness, capability, or both, and specify
                which skills the child is or is not performing).
                {patient.guardian.title_name} reported that {patient.first_name}
                displays issues with sensory processing, including sensory seeking
                behaviors (e.g., ___) and/or sensory sensitivities (e.g., ___ )
                (Include details about the types of input the child seeks or avoids).
            """
            text_color = _RGB.UNRELIABLE
        else:
            text_adaptive = self.llm.run_text_with_parent_input(
                text=f"""
                    Per {patient.guardian.title_name},
                    {patient.first_name} has a history of {PLACEHOLDER}.
                """,
                parent_input=f"""
                    Details on home functioning: {household.home_functioning}
                """,
                comment=(
                    f"{patient.guardian.title_name} reported: "
                    + household.home_functioning
                ),
            )
            text_color = _RGB.LLM

        text_home = string_utils.remove_excess_whitespace(text_home)
        text_adaptive = string_utils.remove_excess_whitespace(text_adaptive)

        self._insert("Home and Adaptive Functioning", word.StyleName.HEADING_2)
        self._insert(text_home)
        self._insert("")
        para = self._insert(text_adaptive)
        cmi_docx.ExtendParagraph(para).format(
            style=cmi_docx.ParagraphStyle(font_rgb=text_color.value)
        )
        self._insert("")

    def write_social_functioning(self) -> None:
        """Writes the social functioning to the report."""
        logger.debug("Writing the social functioning to the report.")
        patient = self.intake.patient
        social_functioning = patient.social_functioning

        self._insert("Social Functioning", word.StyleName.HEADING_2)

        if social_functioning.talents:
            adjectives = self.llm.run_for_adjectives(
                social_functioning.talents,
                comment=f"Talents: {social_functioning.talents}",
            )
            paragraph = self._insert(
                string_utils.remove_excess_whitespace(f"""{patient.guardian.title_name}
                was pleased to describe {patient.first_name} as a {adjectives.lower()}
                {patient.age_gender_label}.""")
            )
        else:
            paragraph = self._insert(
                string_utils.remove_excess_whitespace(f"""{patient.guardian.title_name}
                was pleased to describe {patient.first_name} as a """)
            )
            cmi_docx.ExtendRun(paragraph.add_run(" NO TALENTS PROVIDED ")).format(
                cmi_docx.RunStyle(font_rgb=_RGB.ERROR.value)
            )
            paragraph.add_run(f"{patient.age_gender_label}.")

        llm_text = f"""
            {patient.guardian.title_name} reported that
            {patient.first_name} has {social_functioning.n_friends} close friends
            in {patient.pronouns[2]} peer group.
            {patient.guardian.title_full_name} was concerned about:
            '{social_functioning.social_concerns}'.
            {patient.first_name}'s
            hobbies include {social_functioning.hobbies}.'
        """

        placeholder = self.llm.run_edit(
            text=llm_text,
            additional_instruction="""If no concerns were reported,
                state that the parent reported no concerns. Maintain the
                overall structure of:
                    1. Brief description of child and friends
                    2. Reported concerns.
                    3. Hobbies
                The parent may sometimes provide hobbies in their concerns or
                concerns in the hobbies. Adjust the text accordingly.

                For the number of friends, use a numeric descriptor like
                "no", "few", "several", "many", rather than a specific number.
                As a guideline:
                    0 should be no friends,
                    1-2 should be a few friends,
                    3-4 should be several friends,
                    5+ should be many friends.
            """,
            verify=True,
            comment=string_utils.remove_excess_whitespace(llm_text),
        )

        paragraph.add_run(f" {placeholder}")
        self._insert("")

    def write_psychiatric_history(self) -> None:
        """Writes the psychiatric history to the end of the report."""
        logger.debug("Writing the psychiatric history to the report.")
        self._insert("PSYCHRIATIC HISTORY", word.StyleName.HEADING_1)
        self.write_past_psychiatric_diagnoses()
        self.write_past_psychiatric_hospitalizations()
        self.write_past_therapeutic_interventions()
        self.write_past_psychiatric_medications()
        self.write_past_self_injurious_behaviors_and_suicidality()
        self.write_past_aggressive_behaviors_and_homicidality()
        self.exposure_to_violence_and_trauma()
        self.administration_for_childrens_services_involvement()
        self.write_family_psychiatric_history()

    def write_past_psychiatric_diagnoses(self) -> None:
        """Writes the past psychiatric diagnoses to the report."""
        logger.debug("Writing the past psychiatric diagnoses to the report.")
        patient = self.intake.patient

        past_diagnoses = patient.psychiatric_history.past_diagnoses.base
        if not past_diagnoses:
            text = f"""
                {patient.guardian.title_name} denied any history of past psychiatric
                diagnoses for {patient.first_name}.
            """
            text = string_utils.remove_excess_whitespace(text)
        else:
            instructions = f"""
                The start of the first sentence should be, verbatim, as
                follows: "{patient.guardian.title_name} reported that
                {patient.first_name} was diagnosed with the following psychiatric
                diagnoses: "
            """

            text = self.llm.run_with_object_input(
                items=past_diagnoses,
                additional_instruction=instructions,
                comment="\n\n".join(
                    [str(model) for model in past_diagnoses],
                ),
            )

        self._insert("Past Psychiatric Diagnoses", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_past_psychiatric_hospitalizations(self) -> None:
        """Writes the past psychiatric hospitalizations to the report."""
        logger.debug("Writing the past psychiatric hospitalizations to the report.")
        patient = self.intake.patient
        text = f"""
            {patient.guardian.title_name} denied any history of past psychiatric
            hospitalizations for {patient.first_name}.
        """
        text = string_utils.remove_excess_whitespace(text)

        self._insert("Past Psychiatric Hospitalizations", word.StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(paragraph).format(
            cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        self._insert("")

    def administration_for_childrens_services_involvement(self) -> None:
        """Writes the ACS involvement to the report."""
        logger.debug("Writing the ACS involvement to the report.")

        text = self._write_basic_psychiatric_history(
            "ACS involvement",
            self.intake.patient.psychiatric_history.children_services,
        )
        self._insert(
            "Administration for Children's Services (ACS) Involvement",
            word.StyleName.HEADING_2,
        )
        paragraph = self._insert(text)
        if not self.intake.patient.psychiatric_history.is_follow_up_done:
            cmi_docx.ExtendParagraph(paragraph).format(
                cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
            )
        self._insert("")

    def write_past_aggressive_behaviors_and_homicidality(self) -> None:
        """Writes the past aggressive behaviors and homicidality to the report."""
        logger.debug(
            "Writing the past aggressive behaviors and homicidality to the report.",
        )
        text = self._write_basic_psychiatric_history(
            "homicidality or severe physically aggressive behaviors towards others",
            self.intake.patient.psychiatric_history.aggressive_behaviors,
        )

        self._insert(
            "Past Severe Aggressive Behaviors and Homicidality",
            word.StyleName.HEADING_2,
        )
        paragraph = self._insert(text)
        if not self.intake.patient.psychiatric_history.is_follow_up_done:
            cmi_docx.ExtendParagraph(paragraph).format(
                cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
            )
        self._insert("")

    def write_family_psychiatric_history(self) -> None:
        """Writes the family psychiatric history to the report."""
        text = _FamilyPsychiatricHistory(
            self.intake.patient, self.llm
        ).write_psychiatric_history()

        self._insert("Family Psychiatric History", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_past_therapeutic_interventions(self) -> None:
        """Writes the past therapeutic history to the report."""
        logger.debug("Writing the past therapeutic history to the report.")
        patient = self.intake.patient
        guardian = patient.guardian
        interventions = patient.psychiatric_history.therapeutic_interventions

        if not interventions:
            text = f"""
                    {patient.first_name} does not have any history of mental
                    health treatment.
                """
            text = string_utils.remove_excess_whitespace(text)
        else:
            text = self.llm.run_with_object_input(
                items=interventions,
                additional_instruction=f"""
                    An example structure for this paragraph follows:

                    From {PLACEHOLDER} to {PLACEHOLDER},
                    {patient.first_name} engaged in therapy with
                    {PLACEHOLDER} due to {PLACEHOLDER} at a
                    frequency of {PLACEHOLDER}. {guardian.title_name}
                    described the treatment as {PLACEHOLDER}.
                    Treatment was ended due to {PLACEHOLDER}.
                    """,
                comment="\n\n".join(
                    [str(intervention) for intervention in interventions],
                ),
                verify=True,
            )

        self._insert("Past Therapeutic Interventions", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_past_psychiatric_medications(self) -> None:
        """Writes the past psychiatric medications to the report."""
        logger.debug("Writing the past psychiatric medications to the report.")

        medications = self.intake.patient.psychiatric_history.medications

        self._insert("Past Psychiatric Medications", word.StyleName.HEADING_2)
        if not medications.past_medication and not medications.current_medication:
            text = f"""
                {self.intake.patient.guardian.title_name} denied any history of
                psychiatric medications for {self.intake.patient.first_name}.
            """
            text = string_utils.remove_excess_whitespace(text)
        else:
            if medications.past_medication and medications.current_medication:
                all_medication = (
                    medications.past_medication + medications.current_medication
                )
            elif medications.past_medication:
                all_medication = medications.past_medication  # type: ignore[assignment]
            else:
                all_medication = medications.current_medication  # type: ignore[assignment]

            text = self.llm.run_with_object_input(
                items=all_medication,
                additional_instruction="""
                    You will receive a list of both past and current
                    medications. Please describe this patient's medication
                    history. An example format for past medication follows:

                    "From April 2022 to June 2022, [PATIENT_NAME] was prescribed a
                    course of [MEDICATION] (initial dosage: [INITIAL_DOSAGE],
                    maximum dosage: [MAXIMUM_DOSAGE]) by [DOCTOR NAME].
                    [REASON_FOR_TAKING and RESPONSE_TO_MEDICATION]."

                    For current medications, maintain the same format but change the
                    date format to "Since [DATE_STARTED].
                """,
                comment="\n\n".join([str(model) for model in all_medication]),
            )

        self._insert(text)
        self._insert("")

    def write_past_self_injurious_behaviors_and_suicidality(self) -> None:
        """Writes the past self-injurious behaviors and suicidality to the report."""
        logger.debug(
            "Writing the past self-injurious behaviors and suicidality to the report.",
        )
        text = self._write_basic_psychiatric_history(
            "serious self-injurious harm or suicidal ideation",
            self.intake.patient.psychiatric_history.self_harm,
        )

        self._insert(
            "Past Self-Injurious Behaviors and Suicidality",
            word.StyleName.HEADING_2,
        )
        paragraph = self._insert(text)
        if not self.intake.patient.psychiatric_history.is_follow_up_done:
            cmi_docx.ExtendParagraph(paragraph).format(
                cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
            )
        self._insert("")

    def exposure_to_violence_and_trauma(self) -> None:
        """Writes the exposure to violence and trauma to the report."""
        logger.debug("Writing the exposure to violence and trauma to the report.")

        text = self._write_basic_psychiatric_history(
            "exposure to violence or trauma",
            self.intake.patient.psychiatric_history.violence_and_trauma,
        )

        self._insert("Exposure to Violence and Trauma", word.StyleName.HEADING_2)
        paragraph = self._insert(text)
        if not self.intake.patient.psychiatric_history.is_follow_up_done:
            cmi_docx.ExtendParagraph(paragraph).format(
                cmi_docx.ParagraphStyle(font_rgb=_RGB.UNRELIABLE.value),
            )
        self._insert("")

    def write_medical_history(self) -> None:
        """Writes the medical history to the end of the report."""
        logger.debug("Writing the medical history to the report.")
        patient = self.intake.patient
        primary_care = patient.primary_care

        texts = [
            f"""
            {patient.first_name}'s medical history is unremarkable for
            significant medical conditions. {patient.pronouns[0].capitalize()} is not
            currently taking any medications for chronic medical conditions.""",
            f"""
            {primary_care.glasses_hearing_device}.
            """,
        ]

        prior_diseases = copy.deepcopy(primary_care.prior_diseases)
        if primary_care.head_injury is None:
            prior_diseases.base.append(
                parser_models.PriorDisease(
                    name="head injury",
                    was_positive=False,
                ),
            )
            head_injury_text = ""
        else:
            head_injury_text = self.llm.run_edit(
                primary_care.head_injury,
                additional_instruction=(
                    "The response should be a single sentence that rephrases this "
                    "parent's statement on their child's head injury"
                ),
                context=" ".join(texts),
                comment=primary_care.head_injury,
            )
        texts.append(head_injury_text)
        texts.append(prior_diseases.transform() + ".")
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        self._insert("MEDICAL HISTORY", word.StyleName.HEADING_1)
        paragraph = self._insert(" ".join(texts))
        cmi_docx.ExtendParagraph(paragraph).replace(
            texts[0],
            texts[0],
            cmi_docx.RunStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        self._insert("")

    def write_current_psychiatric_functioning(self) -> None:
        """Writes the current psychiatric functioning to the report."""
        logger.debug("Writing the current psychiatric functioning to the report.")
        self._insert("CURRENT PSYCHIATRIC FUNCTIONING", word.StyleName.HEADING_1)
        self.write_current_psychiatric_medications_intake()
        self.write_current_psychiatric_medications_testing()
        self.write_denied_symptoms()

    def write_current_psychiatric_medications_intake(self) -> None:
        """Writes the current psychiatric medications to the report."""
        logger.debug("Writing the current psychiatric medications to the report.")
        patient = self.intake.patient
        medications = (
            self.intake.patient.psychiatric_history.medications.current_medication
        )

        if not medications:
            text = f"""
                {patient.first_name} is not currently prescribed any psychiatric
                medications.
            """
            text = string_utils.remove_excess_whitespace(text)
        else:
            text = self.llm.run_with_object_input(
                medications,
                comment="\n\n".join(str(med) for med in medications),
                additional_instruction=f"""
                    You will be given a list of current medications taken by a patient.
                    Your job is to write a paragraph appropriate for a clinical report
                    based on this list. Ensure that the phrasing is appropriate for such
                    a report, i.e. if the parent's phrasing is informal or otherwise
                    inappropriate, feel free to change the phrasing without altering
                    the underlying message.

                    The format should be as follows:
                    "{patient.first_name} is current prescribed paracetamol 5mg and
                    Ritalin 5mg for symptoms of hyperactivity.
                    {patient.pronouns[0].capitalize()} is being treated by
                    doctor [NAME], MD. This medication has been effective."

                    If any piece of information is not present in the input data,
                    for example the doctor's name is "don't know", then place a
                    [UNKNOWN NAME] tag where this information should go, where the
                    word NAME may be replaced by the appropriate tag.
                """,
                verify=True,
            )

        self._insert("Current Psychiatric Medications", word.StyleName.HEADING_2)
        self._insert(text)
        self._insert("")

    def write_current_psychiatric_medications_testing(self) -> None:
        """Writes the current psychiatric medications to the report."""
        logger.debug("Writing the current psychiatric medications to the report.")
        patient = self.intake.patient
        texts = [
            f"""
        [Rule out presenting diagnoses, using headlines and KSADS/DSM criteria.
        Examples of headlines include: Temper Outbursts (ending should include
        “{patient.guardian.title_name} denied any consistent patterns of
        irritability for {patient.first_name}" if applicable), Inattention
        and Hyperactivity, Autism-Related Symptoms, Oppositional Defiant
        Behaviors, etc.].""",
            "Establish a baseline first for temper outbursts",
            f"""
           (Ex:
        Though {patient.first_name} is generally a {PLACEHOLDER} child,
        {patient.pronouns[0]} continues to have difficulties with temper
        tantrums…).
        """,
        ]
        texts = [string_utils.remove_excess_whitespace(text) for text in texts]

        paragraph = self._insert(texts[0])
        paragraph.add_run(texts[1])
        paragraph.runs[-1].bold = True
        paragraph.add_run(texts[2])
        cmi_docx.ExtendParagraph(paragraph).format(
            cmi_docx.ParagraphStyle(font_rgb=_RGB.TESTING.value, italic=True),
        )
        self._insert("")

    def write_denied_symptoms(self) -> None:
        """Writes the denied symptoms to the report."""
        logger.debug("Writing the denied symptoms to the report.")
        patient = self.intake.patient
        text = f"""
        {patient.guardian.title_name} and {patient.first_name} denied any
        current significant symptoms related to mood, suicidality, psychosis,
        eating, oppositional or conduct behaviors, substance abuse, autism,
        tics, inattention/hyperactivity, enuresis/encopresis, trauma, sleep,
        panic, anxiety or obsessive-compulsive disorders.
        """
        text = string_utils.remove_excess_whitespace(text)

        header = self._insert("Denied Symptoms", word.StyleName.HEADING_2)
        paragraph = self._insert(text)
        cmi_docx.ExtendParagraph(header).format(
            cmi_docx.ParagraphStyle(font_rgb=_RGB.TESTING.value),
        )
        cmi_docx.ExtendParagraph(paragraph).format(
            cmi_docx.ParagraphStyle(font_rgb=_RGB.TESTING.value),
        )
        self._insert("")

    async def apply_corrections(self) -> None:
        """Applies various grammatical and styling corrections."""
        logger.debug("Applying corrections to the report.")
        document_corrector = language_utils.DocumentCorrections(self.report.document)
        await document_corrector.correct()

    def add_signatures(self) -> None:
        """Adds the signatures to the report.

        Michael Milham's signature is placed in a different location than the
        other signatures. As such, it needs some custom handling.
        """
        logger.debug("Adding signatures to the report.")
        signatures = tuple((DATA_DIR / "signatures").glob("*.png"))

        index = 0
        while index < len(self.report.document.paragraphs):
            text = self.report.document.paragraphs[index].text.lower()
            signature = next(
                (
                    signature
                    for signature in signatures
                    if text.startswith(signature.stem.replace("_", " "))
                ),
                None,
            )

            if signature is None:
                index += 1
            elif signature.stem == "michael_p._milham":
                self.report.insert_image(index, signature)
                index += 2
            else:
                self.report.insert_image(index - 1, signature)
                self.report.insert_paragraph_by_text(index - 1, "")
                index += 3

    def make_llm_edits(self) -> None:
        """Makes edits to the report using a large language model."""
        logger.debug("Making edits to the report using a large language model.")

        for placeholder in self.llm.placeholders:
            self.report.replace(
                placeholder.id,
                placeholder.text,
                cmi_docx.RunStyle(font_rgb=_RGB.LLM.value),
            )
            if placeholder.comment:
                find_runs = self.report.find_in_runs(placeholder.text)
                runs = (
                    (find_runs[-1].runs[0], find_runs[-1].runs[-1])
                    if len(find_runs[-1].runs) > 1
                    else find_runs[-1].runs[0]
                )
                comment.add_comment(
                    self.report.document,
                    runs,
                    author=COMMENT_AUTHOR,
                    text=placeholder.comment,
                )

    def add_footer(self) -> None:
        """Adds a footer to the report."""
        logger.debug("Adding a footer to the report.")
        footer = cmi_docx.ExtendParagraph(
            self.report.document.sections[0].footer.paragraphs[0],
        )
        footer.paragraph.text = (
            "Font Colors: Template, Testing, Large Language Model, Known Error"
        )
        footer.replace(
            "Template",
            "Template",
            cmi_docx.RunStyle(font_rgb=_RGB.UNRELIABLE.value),
        )
        footer.replace(
            "Testing",
            "Testing",
            cmi_docx.RunStyle(font_rgb=_RGB.TESTING.value),
        )
        footer.replace(
            "Large Language Model",
            "Large Language Model",
            cmi_docx.RunStyle(font_rgb=_RGB.LLM.value),
        )
        footer.replace(
            "Known Error",
            "Known Error",
            cmi_docx.RunStyle(font_rgb=_RGB.ERROR.value),
        )

    def add_page_break(self) -> None:
        """Adds a page break to the report."""
        paragraph = self._insert("")
        paragraph.add_run().add_break(enum_text.WD_BREAK.PAGE)

    def _insert(
        self,
        text: str,
        style: word.StyleName = word.StyleName.NORMAL,
    ) -> docx_paragraph.Paragraph:
        """Inserts text at the insertion point.

        Given the current structure of the report, text insertions only occur
        in one location. This point can always be modified by altering the
        `insert_before` attribute.

        Args:
            text: The text to insert.
            style: The style of the text.

        Returns:
            The new paragraph.
        """
        insertion_index = next(
            index
            for index in range(len(self.report.document.paragraphs))
            if self.report.document.paragraphs[index].text == self.insert_before.text
        )
        return self.report.insert_paragraph_by_text(insertion_index, text, style.value)  # type: ignore[no-any-return] # mypy doesn't seem to recognize the return type for an unknown reason.

    def set_superscripts(self) -> None:
        """Sets ordinal suffixes to superscript."""
        logger.debug("Setting ordinal suffixes to superscript.")

        pattern = (
            # Match ordinal numbers with 'th' suffix:
            r"\d*(?:[04-9]|1[1-3])th"  # Matches:
            # - Numbers ending in 0th (10th, 20th, etc.)
            # - Numbers ending in 4-9th (4th, 15th, 26th, etc.)
            # - Numbers ending in 11th, 12th, 13th
            # Match special case ordinals:
            r"|\d*1st"  # Numbers ending in 1st (1st, 21st, 31st, etc.)
            r"|\d*2nd"  # Numbers ending in 2nd (2nd, 22nd, 32nd, etc.)
            r"|\d*3rd"  # Numbers ending in 3rd (3rd, 23rd, 33rd, etc.)
        )

        for paragraph in self.report.document.paragraphs:
            matches = list(re.finditer(pattern, paragraph.text))
            if not matches:
                continue
            for match in matches:
                extended_paragraph = cmi_docx.ExtendParagraph(paragraph)
                suffix = match.string[match.end() - 2 : match.end()]
                run_lengths = list(
                    itertools.accumulate([len(run.text) for run in paragraph.runs])
                )
                run_index = min(
                    bisect.bisect_right(run_lengths, match.start()),
                    len(run_lengths) - 1,
                )
                font_color = paragraph.runs[run_index].font.color
                style = cmi_docx.RunStyle(superscript=True, font_rgb=font_color.rgb)
                extended_paragraph.replace_between(
                    match.end() - 2,
                    match.end(),
                    suffix,
                    style,
                )

    @staticmethod
    def _join_patient_languages(languages: list[parser.Language]) -> str:
        """Joins the patient's languages."""
        fluency_groups = itertools.groupby(
            languages,
            key=lambda language: language.fluency,
        )
        fluency_dict = {
            fluency: [language.name for language in language_group]
            for fluency, language_group in fluency_groups
        }

        language_descriptions = [
            f"{fluency} in {string_utils.oxford_comma(fluency_dict[fluency])}"
            for fluency in ["fluent", "proficient", "conversational"]
            if fluency in fluency_dict
        ]
        prepend_is = len(language_descriptions) > 0
        if "basic" in fluency_dict:
            language_descriptions.append(
                (
                    "has basic skills in "
                    f"{string_utils.oxford_comma(fluency_dict['basic'])}"
                ),
            )

        text = string_utils.oxford_comma(language_descriptions)
        if prepend_is:
            text = "is " + text
        return text

    def _write_basic_psychiatric_history(self, label: str, report: str | None) -> str:
        """Writes the basic template for psychiatric history to the report."""
        if not report:
            text = f"""
                {self.intake.patient.guardian.title_name} denied any history of
                {label} for {self.intake.patient.first_name}.
            """
            return string_utils.remove_excess_whitespace(text)
        return self.llm.run_text_with_parent_input(
            f"{self.intake.patient.guardian.title_name} reported that {PLACEHOLDER}.",
            parent_input=f"Details: {report}",
            comment=report,
        )


class _DisorderMerge(pydantic.BaseModel):
    """Whether diagnoses should be merged in the family psychiatric history."""

    new_name: str
    old_names: tuple[str, ...]
    always_merge: bool


@dataclasses.dataclass(frozen=True)
class _DisorderConfig:
    """Configuration for psychiatric diagnosis display rules."""

    ALWAYS_SHOW: frozenset[str] = frozenset(
        {
            "attention-deficit/hyperactivity disorder",
            "autism",
            "bipolar disorder",
            "conduct disorder",
            "depression",
            "eating disorders",
            "generalized anxiety disorder",
            "obsessive compulsive disorder",
            "oppositional defiant disorder",
            "oppositional defiant or conduct disorders",
            "panic disorder",
            "psychosis",
            "suicidality",
            "specific learning disorders",
            "substance abuse",
        }
    )

    POSITIVE_ONLY: frozenset[str] = frozenset(
        {
            "alcohol abuse",
            "disruptive mood dysregulation disorder",
            "enuresis/encopresis",
            "excoriation",
            "gender dysphoria",
            "intellectual disability",
            "language disorder",
            "personality disorder",
            "PTSD",
            "reactive attachment",
            "selective mutism",
            "social anxiety",
            "separation anxiety",
            "specific phobias",
            "tics/Tourette's",
        }
    )

    @property
    def all_known_diagnoses(self) -> frozenset[str]:
        """Return all known psychiatric diagnoses."""
        return self.ALWAYS_SHOW | self.POSITIVE_ONLY


class _FamilyPsychiatricHistory:
    """Writes the Family Psychiatric History section.

    There's a lot of little rules to this section. The writing of this has been split
    off for clarity.
    """

    def __init__(self, patient: parser.Patient, llm: writer_llm.WriterLlm) -> None:
        self.patient = patient
        self.llm = llm
        self.merge_diagnoses = [
            _DisorderMerge(
                new_name="substance abuse",
                old_names=("alcohol abuse", "substance abuse"),
                always_merge=True,
            ),
            _DisorderMerge(
                new_name="specific learning disorders",
                old_names=(
                    "specific learning disorder, with impairment in mathematics",
                    "specific learning disorder, with impairment in reading",
                    "specific learning disorder, with impairment in written expression",
                ),
                always_merge=True,
            ),
            _DisorderMerge(
                new_name="oppositional defiant or conduct disorders",
                old_names=(
                    "conduct disorder",
                    "oppositional defiant disorder",
                ),
                always_merge=False,
            ),
        ]

    def _merge_diagnoses(
        self, history: list[parser_models.FamilyPsychiatricHistory]
    ) -> list[parser_models.FamilyPsychiatricHistory]:
        """Merges diagnoses where needed.

        Args:
            history: List of the patient's family historical diagnoses.

        Returns:
            Merged diagnoses.
        """
        for merge in self.merge_diagnoses:
            disorders = [
                disorder for disorder in history if disorder.name in merge.old_names
            ]

            if not merge.always_merge:
                is_all_same = all(
                    disorders[0].is_diagnosed == d.is_diagnosed for d in disorders[1:]
                )
                if not is_all_same:
                    continue

            first_degree = any(disorder.first_degree for disorder in disorders)
            second_degree = any(disorder.second_degree for disorder in disorders)
            third_degree = any(disorder.third_degree for disorder in disorders)
            is_diagnosed = any(disorder.is_diagnosed for disorder in disorders)
            family_members = " ".join([disorder.details for disorder in disorders])
            history = [
                disorder for disorder in history if disorder.name not in merge.old_names
            ]

            history.append(
                parser_models.FamilyPsychiatricHistory(
                    name=merge.new_name,
                    first_degree=first_degree,
                    second_degree=second_degree,
                    third_degree=third_degree,
                    is_diagnosed=is_diagnosed,
                    details=family_members,
                )
            )
        return history

    @staticmethod
    def _remove_diagnoses(
        history: list[parser_models.FamilyPsychiatricHistory],
    ) -> list[parser_models.FamilyPsychiatricHistory]:
        """Removes negative diagnoses that need not be reported."""
        for index in range(len(history) - 1, -1, -1):
            disorder = history[index]
            if disorder.name not in _DisorderConfig().all_known_diagnoses:
                msg = f"Unknown disorder: '{disorder.name}'."
                raise KeyError(msg)
            if (
                not disorder.is_diagnosed
                and disorder.name not in _DisorderConfig().ALWAYS_SHOW
            ):
                history.pop(index)
        return history

    def _write_history_unknown(self) -> str:
        return (
            f"{self.patient.possessive_first_name} family psychiatric history "
            f"was unknown."
        )

    def write_psychiatric_history(self) -> str:
        """Write section on the patient's family psychiatric history."""
        history = self.patient.psychiatric_history.family_psychiatric_history
        if not history.is_father_history_known and not history.is_mother_history_known:
            return self._write_history_unknown()

        if not history.is_mother_history_known or not history.is_father_history_known:
            side = "maternal" if history.is_father_history_known else "paternal"
            unknown_history_line = (
                f" Information regarding {self.patient.possessive_first_name} "
                f"{side} family psychiatric history was unknown."
            )
        else:
            unknown_history_line = ""

        diagnoses = self._remove_diagnoses(self._merge_diagnoses(history.diagnoses))
        endorsed_diagnoses = self._get_endorsed_diagnoses(diagnoses)

        denied_diagnoses = [
            diagnosis.name for diagnosis in diagnoses if not diagnosis.is_diagnosed
        ]

        endorsed_sentence = (
            f"{self.patient.first_name}'s family history "
            f"is remarkable for {string_utils.oxford_comma(endorsed_diagnoses)}."
        )
        denied_sentence = (
            f"{self.patient.guardian.title_name} denied any "
            f"family history related to {string_utils.oxford_comma(denied_diagnoses)}."
        )

        first_sentence = (
            (
                f"{self.patient.possessive_first_name} family history is largely "
                "unremarkable for psychiatric illnesses."
            )
            if not endorsed_diagnoses
            else endorsed_sentence
        )

        return first_sentence + " " + denied_sentence + unknown_history_line

    def _get_endorsed_diagnoses(
        self, diagnoses: list[parser_models.FamilyPsychiatricHistory]
    ) -> list[str]:
        """Generates the text for endorsed diagnoses.

        If the parent provided details, use an LLM to summarize. If not, simply
        use the information from the checkboxes (booleans).
        """
        endorsed_diagnoses = []
        for diagnosis in diagnoses:
            if not diagnosis.is_diagnosed:
                continue
            if diagnosis.details:
                family_members = self.llm.classify_family_relatedness(diagnosis.details)
            else:
                # History endorsed, but not specified.
                family_members = string_utils.oxford_comma(
                    [
                        (
                            f"{index + 1}{string_utils.ordinal_suffix(index + 1)} "
                            "degree relative"
                        )
                        for index, diagnosed in enumerate(
                            (
                                diagnosis.first_degree,
                                diagnosis.second_degree,
                                diagnosis.third_degree,
                            )
                        )
                        if diagnosed
                    ]
                )
            endorsed_diagnoses.append(f"{diagnosis.name} ({family_members})")
        return endorsed_diagnoses
