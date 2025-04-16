"""Base class definition for all tables."""

import abc
import enum
import functools
from collections.abc import Callable, Iterable, Mapping, Sequence
from typing import Protocol, Self, TypeVar, runtime_checkable

import cmi_docx
import pydantic
from docx import document, shared, table
from docx.enum import text
from docx.text import paragraph

from ctk_functions.core import config
from ctk_functions.microservices.sql import models

logger = config.get_logger()


class TableDataNotFoundError(Exception):
    """Thrown when a table's data is not found."""


T = TypeVar("T", bound=models.Base)


class ConditionalStyle(pydantic.BaseModel):
    """Applies a style conditional upon a cell's contents.

    Args:
        condition: The condition, if it evaluates to True then the style
            will be applied. Defaults to always True.
        style: The style to apply.
    """

    condition: Callable[[str], bool] = lambda _: True
    style: cmi_docx.CellStyle

    def apply(self, cells: table._Cell | Iterable[table._Cell]) -> None:
        """Applies the style to the cell if the condition is met.

        Args:
            cells: The cells to apply the style to.

        """
        if isinstance(cells, table._Cell):  # noqa: SLF001
            cells = (cells,)

        for cell in cells:
            if self.condition(cell.text):
                cmi_docx.ExtendCell(cell).format(self.style)


class Styles(enum.Enum):
    """Standard styles used throughout the tables."""

    BOLD = ConditionalStyle(
        style=cmi_docx.CellStyle(paragraph=cmi_docx.ParagraphStyle(bold=True))
    )
    LEFT_ALIGN = ConditionalStyle(
        style=cmi_docx.CellStyle(
            paragraph=cmi_docx.ParagraphStyle(
                alignment=text.WD_PARAGRAPH_ALIGNMENT.LEFT
            )
        )
    )
    THICK_TOP_BORDER = ConditionalStyle(
        style=cmi_docx.CellStyle(borders=[cmi_docx.CellBorder(sides=("top",), sz=16)])
    )


class ClinicalRelevance(pydantic.BaseModel):
    """Stores the score ranges for clinical relevance.

    Attributes:
        low: Minimum (exclusive) score for this tier of relevance.
        high: Maximum (inclusive) score for this tier of relevance.
        label: Label for this tier of relevance.
        style: Custom cell styling for this tier of relevance.
        low_inclusive: Whether to include exact equality with the low.
        high_inclusive: Whether to include exact equality with the high.
    """

    low: float | None
    high: float | None
    label: str | None
    style: cmi_docx.CellStyle
    low_inclusive: bool = False
    high_inclusive: bool = True

    def in_range(self, value: float | str) -> bool:
        """Checks if value is within the valid range.

        Excludes low value, includes high value.
        """
        try:
            value = float(value)
        except (ValueError, TypeError):
            # Escape hatch as some tables will mix numeric types with
            # strings to represent missing data.
            logger.debug("Could not check value %s", value)
            return False

        if self.low is not None:
            low_check = self.low.__le__ if self.low_inclusive else self.low.__lt__
        if self.high is not None:
            high_check = self.high.__ge__ if self.high_inclusive else self.high.__gt__

        if self.high is None:
            return low_check(value)
        if self.low is None:
            return high_check(value)
        return low_check(value) and high_check(value)

    def __str__(self) -> str:
        """String representation of the clinical relevance.

        Returns:
            A string denoting the range for this tier's relevance.

        Example outputs:
            "<65 = LABEL" if low is 65, high is not set.
            ">65 = LABEL" if low is not set and high is 65.
            "65-65 = LABEL" if low is 65 and high is 75.
        """
        high_symbol = "<=" if self.high_inclusive else "<"
        low_symbol = ">=" if self.low_inclusive else ">"

        if self.low is None:
            value = f"{high_symbol}{self.high:g}"

        elif self.high is None:
            value = f"{low_symbol}{self.low:g}"
        else:
            value = f"{low_symbol}{self.low:g}, {high_symbol}{self.high:g}"

        if self.label:
            return f"{value} = {self.label}"
        return value

    @pydantic.model_validator(mode="after")
    def _is_low_or_high_defined(self) -> Self:
        if self.high is None and self.low is None:
            msg = "Low or high must be defined."
            raise ValueError(msg)
        return self


class ParagraphBlock(pydantic.BaseModel):
    """A text block in a Word document.

    Attributes:
        content: The text content of the paragraph.
        style: The styling of the paragraph.
        level: The level of the paragraph (e.g. 0 == Title, 1 == Heading 1,
            None == Normal style).
    """

    content: str
    style: cmi_docx.ParagraphStyle = cmi_docx.ParagraphStyle()
    level: int | None = None

    def add_to(self, doc: document.Document) -> paragraph.Paragraph:
        """Adds the paragraph to the end of a document.

        Args:
            doc: The document to add the paragraph to.
        """
        if self.level:
            para = doc.add_heading(text=self.content, level=self.level)
        else:
            para = doc.add_paragraph(self.content)
        if self.style:
            cmi_docx.ExtendParagraph(para).format(self.style)
        return para  # type: ignore[no-any-return]


def default_table_style_factory() -> list[ConditionalStyle]:
    """Creates a list of default table styles.

    Uses a factory method instead of a constant as lists are mutable.

    Returns:
        The default table styles.
    """
    return [
        ConditionalStyle(
            style=cmi_docx.CellStyle(
                paragraph=cmi_docx.ParagraphStyle(
                    space_after=shared.Pt(3),
                    space_before=shared.Pt(3),
                    font_size=11,
                ),
            ),
        )
    ]


class Formatter(pydantic.BaseModel):
    """Formatter for table cells.

    Attributes:
        conditional_styles: A list of styles to apply, conditional on cell contents.
        merge_top: If True, merges cells with identical cells above them.
        merge_right: If True, merges cells with identical cells right of them.
        width: The width of the cell. If None, does not adjust from Word's default.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    conditional_styles: list[ConditionalStyle] = pydantic.Field(
        default_factory=default_table_style_factory,
    )
    merge_top: bool = pydantic.Field(default=False)
    merge_right: bool = pydantic.Field(default=False)
    width: None | shared.Cm | shared.Inches | shared.Pt = None

    def format(self, tbl: table.Table, row_index: int, col_index: int) -> None:
        """Formats the cell content for a table.

        Args:
            tbl: The table to format.
            row_index: The row index of the target cell.
            col_index: The column index of the target cell.
        """
        if row_index == 0 and self.merge_top:
            msg = "Cannot merge top row upwards."
            raise ValueError(msg)
        if col_index == len(tbl.rows[0].cells) and self.merge_right:
            msg = "Cannot merge right row rightwards."
            raise ValueError(msg)

        cell = tbl.columns[col_index].cells[row_index]
        for style in self.conditional_styles:
            style.apply(cell)

        if self.width:
            cell.width = self.width

        if self.merge_top:
            prev_cell = tbl.rows[row_index - 1].cells[col_index]
            current_text = cell.text

            if prev_cell.text == cell.text:
                prev_cell.merge(cell)
                prev_cell.text = current_text

        if self.merge_right:
            next_cell = tbl.rows[row_index].cells[col_index + 1]
            current_text = cell.text
            if next_cell.text == cell.text:
                next_cell.merge(cell)
                cell.text = current_text


class FormatProducer:
    """Producer for table formatting."""

    @classmethod
    def produce(  # noqa: PLR0913
        cls,
        n_rows: int,
        column_widths: Sequence[None | shared.Cm | shared.Inches | shared.Pt],
        global_style: Iterable[ConditionalStyle] | None = None,
        row_styles: Mapping[int, Iterable[ConditionalStyle]] | None = None,
        column_styles: Mapping[int, Iterable[ConditionalStyle]] | None = None,
        cell_styles: Mapping[tuple[int, int], Iterable[ConditionalStyle]] | None = None,
        merge_top: Sequence[int] | None = None,
        *,
        rows_first: bool = True,
    ) -> tuple[tuple[Formatter, ...], ...]:
        """Produces formatters for tables.

        Args:
            n_rows: The number of rows to produce.
            column_widths: The width of the columns.
            global_style: The global style to use, if None, uses the default style
                defined by the default_table_style_factory.
            row_styles: The styles to use for the rows as a mapping
                where the key is the row index and the value is the style.
            column_styles: The styles to use for the columns as a mapping
                where the key is the row index and the value is the style.
            cell_styles: The styles to use for the cells as a mapping
                where the key is the indices (row, column) and the value is the style.
            merge_top: Columns in which to set merge_top to True. Currently only
                supports entire columns, not cell-specific.
            rows_first: If true, applies formatting to rows first, otherwise
                columns first.

        Returns:
            Array of formatters, rows first.
        """
        if global_style is None:
            global_style = default_table_style_factory()

        formatters = tuple(
            [
                tuple(
                    [
                        Formatter(width=width, conditional_styles=global_style)
                        for width in column_widths
                    ]
                )
                for _ in range(n_rows)
            ]
        )

        if rows_first:
            cls._apply_rows(formatters, row_styles)
            cls._apply_columns(formatters, column_styles)
        else:
            cls._apply_columns(formatters, column_styles)
            cls._apply_rows(formatters, row_styles)

        cls._apply_cells(formatters, cell_styles)
        cls._apply_merge_top(formatters, merge_top)
        return formatters

    @staticmethod
    def _apply_cells(
        formatters: Sequence[Sequence[Formatter]],
        styles: Mapping[tuple[int, int], Iterable[ConditionalStyle]] | None,
    ) -> None:
        if not styles:
            return

        for indices, stls in styles.items():
            formatters[indices[0]][indices[1]].conditional_styles.extend(stls)

    @staticmethod
    def _apply_rows(
        formatters: Sequence[Sequence[Formatter]],
        styles: Mapping[int, Iterable[ConditionalStyle]] | None,
    ) -> None:
        if not styles:
            return

        for index, stls in styles.items():
            row = formatters[index]
            for formatter in row:
                formatter.conditional_styles.extend(stls)

    @staticmethod
    def _apply_columns(
        formatters: Sequence[Sequence[Formatter]],
        styles: Mapping[int, Iterable[ConditionalStyle]] | None,
    ) -> None:
        if not styles:
            return

        for index, stls in styles.items():
            for row in formatters:
                row[index].conditional_styles.extend(stls)

    @staticmethod
    def _apply_merge_top(
        formatters: Sequence[Sequence[Formatter]],
        merge_top: Sequence[int] | None,
    ) -> None:
        if not merge_top:
            return

        for index in merge_top:
            for row in formatters[1:]:
                # Top row cannot merge up; skip.
                row[index].merge_top = True


class WordTableCell(pydantic.BaseModel):
    """Definition of a cell in a Word table.

    Attributes:
        content: The contents of the cell.
        formatter: Styling for the cell.
    """

    content: str = pydantic.Field(..., coerce_numbers_to_str=True)
    formatter: Formatter = pydantic.Field(default_factory=Formatter)


class WordTableMarkup(pydantic.BaseModel):
    """Definition of a Word table.

    Attributes:
        rows: The rows of the table.
    """

    model_config = pydantic.ConfigDict(frozen=True)

    rows: tuple[tuple[WordTableCell, ...], ...]

    @pydantic.field_validator("rows")
    @classmethod
    def _rows_are_same_size(
        cls,
        rows: tuple[tuple[WordTableCell, ...], ...],
    ) -> tuple[tuple[WordTableCell, ...], ...]:
        """Checks that the rows form a valid 2D array."""
        if not all(len(row) == len(rows[0]) for row in rows):
            msg = "All rows must have the same length."
            raise ValueError(msg)
        return rows


class DataProducer(abc.ABC):
    """Abstract data producer for Word tables."""

    @classmethod
    @functools.lru_cache
    @abc.abstractmethod
    def fetch(cls, mrn: str) -> tuple[tuple[str, ...], ...]:
        """Abstract data fetcher."""

    @classmethod
    def is_available(cls, mrn: str) -> bool:
        """Tests whether the required data is available."""
        try:
            cls.fetch(mrn)
        except TableDataNotFoundError:
            return False
        return True


class WordDocumentTableRenderer(pydantic.BaseModel):
    """Creates Word tables.

    Attributes:
        markup: The markup for the Word table.
        table_style: The name of the Word table style to use.
    """

    markup: WordTableMarkup
    table_style: str = "Grid Table 7 Colorful"

    def add_to(self, doc: document.Document) -> None:
        """Adds the table to the document.

        Args:
            doc: The document to add the table to.
        """
        n_rows = len(self.markup.rows)
        n_cols = len(self.markup.rows[0])
        tbl = doc.add_table(n_rows, n_cols)
        tbl.style = self.table_style

        for row_index in range(n_rows):
            for col_index in range(n_cols):
                document_cell = tbl.rows[row_index].cells[col_index]
                template_cell = self.markup.rows[row_index][col_index]

                document_cell.text = template_cell.content

        for col_index in range(n_cols):
            for row_index in range(n_rows):
                # Formatting must be done after all content is added as
                # adding more content may conflict with previously set
                # cell widths.
                template_cell = self.markup.rows[row_index][col_index]
                template_cell.formatter.format(tbl, row_index, col_index)


class WordDocumentTableSectionRenderer(pydantic.BaseModel):
    """Creates a section around a Word table.

    Many of the Pyrite report tables have inconsistent paragraphs around them.
    This class is used to handle all these cases.

    Attributes:
        preamble: Text content to appear before the paragraph.
        table_renderer: The table renderer.
        postamble: Text content to appear after the paragraph. Note: if no paragraph
            appears between two tables, then the tables will be merged.
    """

    table_renderer: WordDocumentTableRenderer
    preamble: Sequence[ParagraphBlock] = pydantic.Field(default_factory=list)
    postamble: Sequence[ParagraphBlock] = [ParagraphBlock(content="")]

    def add_to(self, doc: document.Document) -> None:
        """Adds the section to the document.

        Args:
            doc: The document to add the section to.
        """
        for pre in self.preamble:
            pre.add_to(doc)

        self.table_renderer.add_to(doc)

        for post in self.postamble:
            post.add_to(doc)


class WordTableSection(abc.ABC):
    """Abstract class for adding table sections."""

    data_source: type[DataProducer]
    mrn: str

    @abc.abstractmethod
    def __init__(self, mrn: str) -> None:
        """Initialize the section for a given MRN.

        Args:
            mrn: The participant's unique identifier.
        """

    @abc.abstractmethod
    def add_to(self, doc: document.Document) -> None:
        """Adds the section to the document.

        Args:
            doc: The document to add the section to.
        """

    def is_available(self) -> bool:
        """Convenience method for the data source's availability."""
        return self.data_source.is_available(self.mrn)


@runtime_checkable
class _AddToProtocol(Protocol):
    """Protocol defining what's required to use the WordTableSectionAddToMixin."""

    @property
    def mrn(self) -> str:
        """The participant's unique identifier."""

    @property
    def data_source(self) -> DataProducer:
        """The data source to use."""

    @property
    def formatters(self) -> Sequence[Sequence[Formatter]]:
        """The formatters to apply to the data."""


class WordTableSectionAddToMixin:
    """Adds a standardized add_to method for WordTableSections.

    Most tables' add_to functions follow a standardized pattern. This mixin
    provides the option of adding the standard pattern.
    """

    def add_to(self, doc: document.Document) -> None:
        """Adds the data source to the document.

        Args:
            doc: The document to add the section to.
        """
        if not isinstance(self, _AddToProtocol):
            msg = (
                "Classes using the AddToMixin must be a valid implementation of "
                "the AddToProtocol."
            )
            raise TypeError(msg)

        text = self.data_source.fetch(self.mrn)
        rows = [
            [
                WordTableCell(content=text, formatter=formatter)
                for text, formatter in zip(text_row, formatter_row, strict=True)
            ]
            for text_row, formatter_row in zip(text, self.formatters, strict=True)
        ]
        markup = WordTableMarkup(rows=rows)

        args = {
            "preamble": getattr(self, "preamble", None),
            "postamble": getattr(self, "postamble", None),
            "table_renderer": WordDocumentTableRenderer(markup=markup),
        }
        args = {key: val for key, val in args.items() if val}

        renderer = WordDocumentTableSectionRenderer(**args)
        renderer.add_to(doc)
