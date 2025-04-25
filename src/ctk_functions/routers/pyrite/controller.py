"""Business logic for the Pyrite endpoints."""

import io
from typing import Any

import cmi_docx
import docx
import fastapi
from docx.text import paragraph as docx_paragraph
from fastapi import status

from ctk_functions.core import config
from ctk_functions.microservices.sql import models
from ctk_functions.routers.pyrite import sql_data
from ctk_functions.routers.pyrite.reports import reports
from ctk_functions.routers.pyrite.tables import (
    base,
)

logger = config.get_logger()
settings = config.get_settings()

DATA_DIR = settings.DATA_DIR


def get_pyrite_report(mrn: str) -> bytes:
    """Generates a Pyrite report for a given MRN.

    Args:
        mrn: The participant's identifier.

    Returns:
        The .docx file bytes.
    """
    logger.debug("Entered controller of get_pyrite_report.")
    report = PyriteReport(mrn)
    report.create(version="alabaster")

    logger.debug("Successfully generated Pyrite report.")
    out = io.BytesIO()
    report.document.save(out)
    return out.getvalue()


class PyriteReport:
    """Builder of the Pyrite reports.

    Pyrite reports contain the overall test
    """

    def __init__(self, mrn: str) -> None:
        """Initialize the Pyrite report.

        Args:
            mrn: The participant's unique identifier.
        """
        self._mrn = mrn
        self.document = docx.Document(str(DATA_DIR / "pyrite_template.docx"))

    def create(self, version: reports.VERSIONS, **kwargs: Any) -> None:  # noqa: ANN401
        """Creates the Pyrite report.

        Args:
            version: The version of the report to generate.
            **kwargs: Version-specific keyword arguments.
        """
        structure = reports.get_report_structure(self._mrn, version, **kwargs)
        for section in structure:
            section.add_to(self.document)

        # As an artifact from using a template file, the first paragraph is empty.
        # Delete it.
        self._delete_paragraph(self.document.paragraphs[0])

        self._replace_participant_information()

    def _get_participant(self) -> models.CmiHbnIdTrack:
        """Fetches the participant's data from the SQL database.

        Returns:
            A row from the CMI_HB_IDTrack_t table.
        """
        sanitized_mrn = self._mrn.replace("\r", "").replace("\n", "")
        logger.debug("Fetching participant %s.", sanitized_mrn)
        try:
            return sql_data.fetch_participant_row(
                "MRN", self._mrn, models.CmiHbnIdTrack
            )  # type: ignore[no-any-return, unused-ignore] # Getting errors both when no-any-return is, and is not used.
        except base.TableDataNotFoundError as exception_info:
            raise fastapi.HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="MRN not found.",
            ) from exception_info

    def _replace_participant_information(self) -> None:
        """Replaces the patient information in the report."""
        logger.debug("Replacing patient information in the report.")
        participant = self._get_participant()
        first_name = participant.first_name
        full_name = f"{first_name} {participant.last_name}"
        first_name_possessive = (
            f"{first_name}{"'" if first_name.endswith('s') else "'s"}"
        )
        replacements = {
            "full_name": full_name,
            "first_name_possessive": first_name_possessive,
        }

        for template, replacement in replacements.items():
            template_formatted = "{{" + template.upper() + "}}"
            cmi_docx.ExtendDocument(self.document).replace(
                template_formatted,
                replacement,
            )

    @staticmethod
    def _delete_paragraph(para: docx_paragraph.Paragraph) -> None:
        """Deletes a paragraph.

        Args:
            para: The paragraph to delete.
        """
        p_elem = para._element  # noqa: SLF001
        p_elem.getparent().remove(p_elem)
        p_elem._p = p_elem._element = None  # noqa: SLF001
