"""Contains definitions of Pyrite report formats."""

import abc
import dataclasses
from collections.abc import Callable
from typing import Literal

import cmi_docx
import pydantic
from docx import document
from docx.enum import text as enum_text

from ctk_functions.routers.pyrite.tables import (
    academic_achievement,
    base,
    cbcl_ysr,
    celf5,
    conners3,
    ctopp2,
    grooved_pegboard,
    language,
    mfq,
    scared,
    scq,
    srs,
    swan,
    wisc_composite,
    wisc_subtest,
)

VERSIONS = Literal["alabaster"]


class Section(abc.ABC):
    """Represents a section in the report structure."""

    @abc.abstractmethod
    def add_to(self, doc: document.Document) -> None:
        """Adds the section to the report."""


class PageBreak(Section):
    """A page break in the report."""

    def add_to(self, doc: document.Document) -> None:
        """Adds a page break to the report.

        Args:
            doc: The document to add the page break to.
        """
        doc.add_paragraph().add_run().add_break(enum_text.WD_BREAK.PAGE)


VALID_STYLES = Literal["Heading 1", "Heading 2", "Heading 3", "Heading 1 Centered"]


class ParagraphSection(pydantic.BaseModel, Section):
    """Represents a text block in the report structure.

    Attributes:
        content: The textual content of the paragraph.
        style: Either a ParagraphStyle defining the styling, a string
            matching a Style name in the document, or None. If None,
            uses the document's default style.
        condition: Condition to evaluate for section inclusion.
        subsections: Subsections of this section, will be appended if
            this section's condition and the subsection's condition is met.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    content: str
    style: cmi_docx.ParagraphStyle | VALID_STYLES | None = None
    condition: Callable[[], bool] = lambda: True
    subsections: list[Section] = pydantic.Field(default_factory=list)

    def add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if isinstance(self.style, cmi_docx.ParagraphStyle):
            para = doc.add_paragraph(self.content)
            extended_paragraph = cmi_docx.ExtendParagraph(para)
            extended_paragraph.format(self.style)
        else:
            doc.add_paragraph(self.content, self.style)
        for subsection in self.subsections:
            subsection.add_to(doc)


class TableSection(pydantic.BaseModel, Section):
    """Represents a table section in the report structure."""

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    title: str | None = None
    level: int | None = None
    condition: Callable[[], bool] = lambda: True
    tables: list[base.WordTableSection] = pydantic.Field(default_factory=list)
    subsections: list[Section] = pydantic.Field(default_factory=list)

    def add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if not self.condition():
            return
        if self.title and self.level:
            doc.add_heading(self.title, level=self.level)
        for table in self.tables:
            if table.is_available():
                table.add_to(doc)
        for subsection in self.subsections:
            subsection.add_to(doc)


@dataclasses.dataclass
class _PyriteTableCollection:
    """Collection of all tables used in Pyrite reports."""

    def __init__(self, mrn: str) -> None:
        """Initializes all tables."""
        self.wisc_composite = wisc_composite.WiscCompositeTable(mrn)
        self.wisc_subtest = wisc_subtest.WiscSubtestTable(mrn)
        self.grooved_pegboard = grooved_pegboard.GroovedPegboardTable(mrn)
        self.academic_achievement = academic_achievement.AcademicAchievementTable(mrn)
        self.celf5 = celf5.Celf5Table(mrn)
        self.language = language.LanguageTable(mrn)
        self.ctopp2 = ctopp2.Ctopp2Table(mrn)
        self.cbcl = cbcl_ysr.CbclTable(mrn)
        self.ysr = cbcl_ysr.YsrTable(mrn)
        self.swan = swan.SwanTable(mrn)
        self.conners3 = conners3.Conners3Table(mrn)
        self.scq = scq.ScqTable(mrn)
        self.srs = srs.SrsTable(mrn)
        self.mfq = mfq.MfqTable(mrn)
        self.scared = scared.ScaredTable(mrn)


def get_report_structure(
    mrn: str,
    version: VERSIONS,
) -> tuple[Section, ...]:
    """Fetches a report structure based on version name.

    Valid versions are:
        alabaster: Maintained from 2025-04-02 until the present. This version
            outputs all available tables and is currently the 'main' version.
        chalk: Debug version that runs requested tables only.

    Args:
        mrn: The participant's unique identifier.
        version: The report version name.

    Returns:
          The structure of the Pyrite report.
    """
    if version == "alabaster":
        return _report_alabaster(mrn)
    msg = f"Invalid Pyrite version: {version}."
    raise ValueError(msg)


def _report_alabaster(mrn: str) -> tuple[Section, ...]:
    """Creates the structure of the 2024-04-02 Pyrite report.

    Args:
        mrn: The participant's unique identifier.

    Returns:
        The structure of the Pyrite report, containing only available sections.
    """

    def is_any_available(*tbls: base.WordTableSection) -> bool:
        return any(tbl.is_available() for tbl in tbls)

    def is_all_available(*tbls: base.WordTableSection) -> bool:
        return all(tbl.is_available() for tbl in tbls)

    tables = _PyriteTableCollection(mrn)

    return (
        PageBreak(),  # Start with a page break from the end of the template file.
        ParagraphSection(content="Results Appendix", style="Heading 1 Centered"),
        ParagraphSection(
            content="General Intellectual Function",
            style="Heading 1",
            condition=lambda: is_all_available(
                tables.wisc_composite, tables.wisc_subtest
            ),
            subsections=[
                TableSection(
                    title=(
                        "The Wechsler Intelligence Scale for Children-Fifth "
                        "Edition (WISC-V)"
                    ),
                    level=3,
                    tables=[tables.wisc_composite, tables.wisc_subtest],
                ),
            ],
        ),
        ParagraphSection(
            content="Fine Motor Dexterity",
            style="Heading 1",
            condition=lambda: tables.grooved_pegboard.is_available(),
            subsections=[
                TableSection(
                    title="Lafayette Grooved Pegboard Test",
                    level=3,
                    tables=[tables.grooved_pegboard],
                )
            ],
        ),
        PageBreak(),
        TableSection(
            title="Academic Achievement",
            level=1,
            tables=[tables.academic_achievement],
            condition=lambda: tables.academic_achievement.is_available(),
        ),
        PageBreak(),
        TableSection(
            title="Language Skills",
            level=1,
            tables=[tables.celf5, tables.language, tables.ctopp2],
            condition=lambda: is_any_available(
                tables.celf5,
                tables.language,
                tables.ctopp2,
            ),
        ),
        PageBreak(),
        ParagraphSection(
            content="Social-Emotional and Behavioral Functioning Questionnaires",
            style="Heading 1",
            condition=lambda: is_any_available(
                tables.cbcl,
                tables.ysr,
                tables.swan,
                tables.conners3,
                tables.scq,
                tables.srs,
                tables.mfq,
                tables.scared,
            ),
            subsections=[
                TableSection(
                    title="Child Behavior Checklist - Parent Report Form (CBCL)",
                    level=3,
                    tables=[tables.cbcl],
                    condition=lambda: tables.cbcl.is_available(),
                ),
                TableSection(
                    title="Child Behavior Checklist - Youth Self Report (YSR)",
                    level=3,
                    tables=[tables.ysr],
                    condition=lambda: tables.ysr.is_available(),
                ),
                ParagraphSection(
                    content="Attention Deficit-Hyperactivity Symptoms and Behaviors",
                    style="Heading 1",
                    condition=lambda: is_any_available(
                        tables.swan,
                        tables.conners3,
                    ),
                    subsections=[
                        TableSection(
                            title=(
                                "Strengths and Weaknesses of ADHD Symptoms "
                                "and Normal Behavior (SWAN)"
                            ),
                            level=3,
                            tables=[tables.swan],
                            condition=lambda: tables.swan.is_available(),
                        ),
                        TableSection(
                            title="Conners 3 - Child Short Form",
                            level=3,
                            tables=[tables.conners3],
                            condition=lambda: tables.conners3.is_available(),
                        ),
                    ],
                ),
                ParagraphSection(
                    content="Autism Spectrum Symptoms and Behaviors",
                    style="Heading 1",
                    condition=lambda: is_any_available(
                        tables.scq,
                        tables.srs,
                    ),
                    subsections=[
                        TableSection(
                            title="Social Communication Questionnaire",
                            level=3,
                            tables=[tables.scq],
                            condition=lambda: tables.scq.is_available(),
                        ),
                        TableSection(
                            title="Social Responsiveness Scale",
                            level=3,
                            tables=[tables.srs],
                            condition=lambda: tables.srs.is_available(),
                        ),
                    ],
                ),
                PageBreak(),
                ParagraphSection(
                    content="Depression and Anxiety Symptoms",
                    style="Heading 1",
                    condition=lambda: is_any_available(
                        tables.mfq,
                        tables.scared,
                    ),
                    subsections=[
                        TableSection(
                            title=(
                                "Mood and Feelings Questionnaire (MFQ) - Long Version"
                            ),
                            level=3,
                            tables=[tables.mfq],
                            condition=lambda: tables.mfq.is_available(),
                        ),
                        TableSection(
                            title="Screen for Child Anxiety Related Disorders",
                            level=3,
                            tables=[tables.scared],
                            condition=lambda: tables.scared.is_available(),
                        ),
                    ],
                ),
            ],
        ),
    )
