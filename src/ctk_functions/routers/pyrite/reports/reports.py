"""Contains definitions of Pyrite report formats."""

import abc
import enum
from collections.abc import Callable, Iterable
from typing import Literal, TypeVar

import cmi_docx
import pydantic
from docx import document
from docx.enum import text as enum_text

from ctk_functions.core import config
from ctk_functions.routers.pyrite.reports import appendix_a, introduction, utils
from ctk_functions.routers.pyrite.tables import (
    academic_achievement,
    base,
    cbc,
    celf5,
    conners3,
    ctopp2,
    grooved_pegboard,
    language,
    mfq,
    scared,
    scq,
    srs,
    swan,
    wisc_composite,
    wisc_subtest,
)
from ctk_functions.routers.pyrite.tables import utils as tables_utils

settings = config.get_settings()
DATA_DIR = settings.DATA_DIR
VERSIONS = Literal["alabaster"]
VALID_PARAGRAPH_STYLES = Literal[
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 1 Centered",
    "Normal",
    "Normal Hanging",  # After first line, 0.7cm indent.
]
T = TypeVar("T")
U = TypeVar("U")


class RunStyles(enum.StrEnum):
    """Valid styles in the Word template for runs."""

    Emphasis = "Emphasis"


class Section(pydantic.BaseModel, abc.ABC):
    """Represents a section in the report structure."""

    subsections: list["Section"] = pydantic.Field(default_factory=list)
    condition: Callable[[], bool] = lambda: True

    def add_to(self, doc: document.Document) -> None:
        """Adds the section to the report."""
        if not self.condition():
            return

        self._add_to(doc)
        for subsection in self.subsections:
            subsection.add_to(doc)

    @abc.abstractmethod
    def _add_to(self, doc: document.Document) -> None:
        """Adds this section to the report.

        Handling of the conditional and subsections is done by self.add_to()
        """


class PageBreak(Section):
    """A page break in the report."""

    def _add_to(self, doc: document.Document) -> None:
        """Adds a page break to the report.

        If any paragraph exists, appends it to the last paragraph. Otherwise,
        create

        Args:
            doc: The document to add the page break to.
        """
        para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
        para.add_run().add_break(enum_text.WD_BREAK.PAGE)


class RunsSection(Section):
    """Represents a text block in the report structure with sub-formatting.

    Attributes:
        content: The textual content of each run.
        run_styles: The style of each run. Can be a valid CharacterStyle,
            a RunStyle object, or None (document's default style).
        condition: Condition to evaluate for section inclusion.
        subsections: Subsections of this section, will be appended if
            this section's condition and the subsection's condition is met.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    content: tuple[str, ...]
    paragraph_style: cmi_docx.ParagraphStyle | VALID_PARAGRAPH_STYLES | None = None
    run_styles: tuple[cmi_docx.RunStyle | RunStyles | None, ...]

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if isinstance(self.paragraph_style, cmi_docx.ParagraphStyle):
            para = doc.add_paragraph("")
            extended_paragraph = cmi_docx.ExtendParagraph(para)
            extended_paragraph.format(self.paragraph_style)
        else:
            para = doc.add_paragraph("", self.paragraph_style)

        for text, style in zip(self.content, self.run_styles, strict=True):
            if style is None or isinstance(style, str):
                para.add_run(text, style)
                continue

            run = para.add_run(text)
            extend_run = cmi_docx.ExtendRun(run)
            extend_run.format(style)


class ParagraphSection(Section):
    """Represents a text block in the report structure.

    Attributes:
        content: The textual content of the paragraph.
        style: Either a ParagraphStyle defining the styling, a string
            matching a Style name in the document, or None. If None,
            uses the document's default style.
        condition: Condition to evaluate for section inclusion.
        subsections: Subsections of this section, will be appended if
            this section's condition and the subsection's condition is met.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    content: str
    style: cmi_docx.ParagraphStyle | VALID_PARAGRAPH_STYLES | None = None

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if isinstance(self.style, cmi_docx.ParagraphStyle):
            para = doc.add_paragraph(self.content)
            extended_paragraph = cmi_docx.ExtendParagraph(para)
            extended_paragraph.format(self.style)
        else:
            doc.add_paragraph(self.content, self.style)


class TableSection(Section):
    """Represents a table section in the report structure."""

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    title: str | None = None
    level: int | None = None
    tables: list[base.WordTableSection] = pydantic.Field(default_factory=list)

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if self.title and self.level:
            doc.add_heading(self.title, level=self.level)
        for table in self.tables:
            if table.is_available():
                table.add_to(doc)


class ImageSection(Section):
    """Represents a image section in the report structure."""

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)
    path: pydantic.FilePath

    def _add_to(self, doc: document.Document) -> None:
        """Adds the image to a document.

        Args:
            doc: The document to add the image to.
        """
        doc.add_picture(str(self.path))


class _PyriteTableCollection:
    """Collection of all tables used in Pyrite reports."""

    def __init__(self, mrn: str) -> None:
        """Initializes all tables."""
        self.asr = cbc.AsrTable(mrn)
        self.academic_achievement = academic_achievement.AcademicAchievementTable(mrn)
        self.cbcl = cbc.CbclTable(mrn)
        self.celf5 = celf5.Celf5Table(mrn)
        self.conners3 = conners3.Conners3Table(mrn)
        self.ctopp2 = ctopp2.Ctopp2Table(mrn)
        self.grooved_pegboard = grooved_pegboard.GroovedPegboardTable(mrn)
        self.language = language.LanguageTable(mrn)
        self.mfq = mfq.MfqTable(mrn)
        self.scared = scared.ScaredTable(mrn)
        self.scq = scq.ScqTable(mrn)
        self.srs = srs.SrsTable(mrn)
        self.swan = swan.SwanTable(mrn)
        self.trf = cbc.TrfTable(mrn)
        self.wisc_composite = wisc_composite.WiscCompositeTable(mrn)
        self.wisc_subtest = wisc_subtest.WiscSubtestTable(mrn)
        self.ysr = cbc.YsrTable(mrn)


def get_report_structure(
    mrn: str,
    version: VERSIONS,
) -> tuple[Section, ...]:
    """Fetches a report structure based on version name.

    Valid versions are:
        alabaster: Maintained from 2025-04-02 until the present. This version
            outputs all available tables and is currently the 'main' version.
        chalk: Debug version that runs requested tables only.

    Args:
        mrn: The participant's unique identifier.
        version: The report version name.

    Returns:
          The structure of the Pyrite report.
    """
    if version == "alabaster":
        return _report_alabaster(mrn)
    msg = f"Invalid Pyrite version: {version}."
    raise ValueError(msg)


def _report_alabaster(mrn: str) -> tuple[Section, ...]:
    """Creates the structure of the 2024-04-02 Pyrite report.

    Args:
        mrn: The participant's unique identifier.

    Returns:
        The structure of the Pyrite report, containing only available sections.
    """
    tables = _PyriteTableCollection(mrn)
    table_sections = _get_alabaster_table_structure(tables)
    intro = _table_sections_to_introduction(mrn, table_sections)
    appendix = _table_sections_to_appendix_a(mrn, table_sections)
    return *intro, PageBreak(), *appendix, PageBreak(), *table_sections


def _table_sections_to_introduction(
    mrn: str, table_sections: Iterable[Section]
) -> tuple[Section, ...]:
    """Creates the introduction section of the report.

    Args:
        mrn: The participant's unique identifier.
        table_sections: The list of sections included in the tables.

    Returns:
        The introduction section of the report.
    """
    unique_test_ids = _tables_to_test_ids(mrn, table_sections)
    overviews = introduction.TestOverviewManager()
    used_overviews = [overviews.fetch(test_id) for test_id in unique_test_ids]
    sections = [
        _description_to_overview(mrn, overview)
        for overview in used_overviews
        if overview
    ]
    return (
        ParagraphSection(
            content="STANDARDIZED TESTING, INTERVIEW AND QUESTIONNAIRE RESULTS",
            style="Heading 1",
        ),
        ParagraphSection(
            content=(
                "This report provides a summary of the following tests, "
                "questionnaires, and clinical interviews that were administered during "
                "participation in the study. Areas assessed include general cognitive "
                "ability, academic achievement, language and fine motor coordination. "
                "Clinical interviews and questionnaires were used to assess social, "
                "emotional and behavioral functioning. Please reference Appendix A for "
                "a full description of the measures administered in the Healthy Brain "
                "Network research protocol."
            ),
            style=None,
        ),
        *sections,
        PageBreak(),
        ParagraphSection(
            content="What do the scores represent?",
            style="Heading 2",
        ),
        ParagraphSection(
            content=(
                "Standard scores, T scores, and percentile ranks can indicate "
                "{{FIRST_NAME_POSSESSIVE}} performance compared to other children in "
                "the same age or same grade (see Figure below). A standard score of "
                "100 is the mean of the normative sample (individuals in the same age "
                "or same grade), and a standard score within the range of 90-109 "
                "indicates that an individual's performance or rating is within the "
                "average or typical range. A T score of 50 is the mean of the "
                "normative sample, and a T score within the range of 43-57 indicates "
                "that an individual's performance or ratings is within the average "
                "range. A percentile rank of 50 is the median of the normative sample, "
                "meaning that an individual's performance equals or exceeds 50% of the "
                "same-age peers. A percentile rank within the range of 25-75 indicates "
                "that an individual's performance or rating is within the average."
            ),
            style=None,
        ),
        ImageSection(path=DATA_DIR / "pyrite_tscore_distribution.png"),
    )


def _table_sections_to_appendix_a(
    mrn: str, table_sections: Iterable[Section]
) -> tuple[Section, ...]:
    """Extracts all used DataProducers and converts this information to Appendix A.

    Args:
        mrn: The participant's unique identifier.
        table_sections: The list of sections to extract producers from.

    Returns:
        The section structure for Appendix A.
    """
    unique_test_ids = _tables_to_test_ids(mrn, table_sections)
    descriptions = appendix_a.TestDescriptionManager()
    used_descriptions = [descriptions.fetch(test_id) for test_id in unique_test_ids]
    sections = [_description_to_section(desc) for desc in used_descriptions]
    return (
        ParagraphSection(
            content="Appendix A. Instruments administered in Healthy Brain Network",
            style="Heading 1",
        ),
        ParagraphSection(content=""),
        *sections,
    )


def _tables_to_test_ids(
    mrn: str, table_sections: Iterable[Section]
) -> list[utils.TestId]:
    """Converts tables to the test_ids that they use.

    Args:
        mrn: The participant's unique identifier.
        table_sections: The list of sections to extract IDs from.

    Returns:
        Unique test IDs sorted alphabetically.
    """
    used_producers = _extract_producers_used(table_sections)
    available_producers = [
        producer for producer in used_producers if producer.is_available(mrn)
    ]
    test_ids: list[utils.TestId] = _flatten(
        [tbl.test_ids(mrn) for tbl in available_producers]
    )
    return sorted(dict.fromkeys(test_ids))


def _description_to_section(description: appendix_a.TestDescription) -> Section:
    """Converts a test description to a section."""
    return ParagraphSection(
        content=description.title,
        style="Heading 2",
        subsections=[
            ParagraphSection(
                content=description.description,
            ),
            RunsSection(
                content=("Reference:", " " + description.reference),
                run_styles=(RunStyles.Emphasis, None),
            ),
        ],
    )


def _description_to_overview(mrn: str, overview: introduction.TestOverview) -> Section:
    """Converts a test description to a section."""
    if overview.date_table:
        data = tables_utils.fetch_participant_row(
            "person_id", mrn, overview.date_table.model
        )
        administer_date = str(getattr(data, overview.date_table.column))
    else:
        administer_date = ""

    if not overview.description:
        texts: tuple[str, ...] = (overview.title + "\n", "Administered on: ")
        run_styles: tuple[None | RunStyles, ...] = (None, None)
    else:
        texts = (
            overview.title + "\n",
            overview.description + "\n",
            f"Administered on: {administer_date}",
        )
        run_styles = (None, RunStyles.Emphasis, None)

    return RunsSection(
        content=texts,
        paragraph_style="Normal Hanging",
        run_styles=run_styles,
    )


def _extract_producers_used(
    structure: Iterable[Section],
) -> tuple[type[base.DataProducer], ...]:
    """Extracts the data producers used in sections."""
    producers = []
    for section in structure:
        section_tables: list[base.WordTableSection] = getattr(section, "tables", [])
        producers.extend([tbl.data_source for tbl in section_tables])
        producers.extend(_extract_producers_used(section.subsections))
    return tuple(producers)


def _get_alabaster_table_structure(
    tables: _PyriteTableCollection,
) -> tuple[Section, ...]:
    """Defines the structure of the Alabaster tables."""
    return (
        ParagraphSection(content="Results Appendix", style="Heading 1 Centered"),
        ParagraphSection(
            content="General Intellectual Function",
            style="Heading 1",
            condition=lambda: _all_tables_available(
                tables.wisc_composite, tables.wisc_subtest
            ),
            subsections=[
                TableSection(
                    title=(
                        "The Wechsler Intelligence Scale for Children-Fifth "
                        "Edition (WISC-V)"
                    ),
                    level=3,
                    tables=[tables.wisc_composite, tables.wisc_subtest],
                ),
            ],
        ),
        ParagraphSection(
            content="Fine Motor Dexterity",
            style="Heading 1",
            condition=lambda: tables.grooved_pegboard.is_available(),
            subsections=[
                TableSection(
                    title="Lafayette Grooved Pegboard Test",
                    level=3,
                    tables=[tables.grooved_pegboard],
                )
            ],
        ),
        PageBreak(),
        TableSection(
            title="Academic Achievement",
            level=1,
            tables=[tables.academic_achievement],
            condition=lambda: tables.academic_achievement.is_available(),
        ),
        PageBreak(),
        TableSection(
            title="Language Skills",
            level=1,
            tables=[tables.celf5, tables.language, tables.ctopp2],
            condition=lambda: _any_table_available(
                tables.celf5,
                tables.language,
                tables.ctopp2,
            ),
        ),
        PageBreak(),
        ParagraphSection(
            content="Social-Emotional and Behavioral Functioning Questionnaires",
            style="Heading 1 Centered",
            condition=lambda: _any_table_available(
                tables.cbcl,
                tables.ysr,
                tables.swan,
                tables.conners3,
                tables.scq,
                tables.srs,
                tables.mfq,
                tables.scared,
            ),
            subsections=[
                TableSection(
                    title="Child Behavior Checklist - Parent Report Form (CBCL)",
                    level=3,
                    tables=[tables.cbcl],
                    condition=lambda: tables.cbcl.is_available(),
                ),
                TableSection(
                    title="Child Behavior Checklist - Youth Self Report (YSR)",
                    level=3,
                    tables=[tables.ysr],
                    condition=lambda: tables.ysr.is_available(),
                ),
                TableSection(
                    title="Child Behavior Checklist - Teacher Report Form (TRF)",
                    level=3,
                    tables=[tables.trf],
                    condition=lambda: tables.trf.is_available(),
                ),
                TableSection(
                    title="Child Behavior Checklist - Adult Self Report Form (ASR)",
                    level=3,
                    tables=[tables.asr],
                    condition=lambda: tables.asr.is_available(),
                ),
                ParagraphSection(
                    content="Attention Deficit-Hyperactivity Symptoms and Behaviors",
                    style="Heading 1",
                    condition=lambda: _any_table_available(
                        tables.swan,
                        tables.conners3,
                    ),
                    subsections=[
                        TableSection(
                            title=(
                                "Strengths and Weaknesses of ADHD Symptoms "
                                "and Normal Behavior (SWAN)"
                            ),
                            level=3,
                            tables=[tables.swan],
                            condition=lambda: tables.swan.is_available(),
                        ),
                        TableSection(
                            title="Conners 3 - Child Short Form",
                            level=3,
                            tables=[tables.conners3],
                            condition=lambda: tables.conners3.is_available(),
                        ),
                    ],
                ),
                ParagraphSection(
                    content="Autism Spectrum Symptoms and Behaviors",
                    style="Heading 1",
                    condition=lambda: _any_table_available(
                        tables.scq,
                        tables.srs,
                    ),
                    subsections=[
                        TableSection(
                            title="Social Communication Questionnaire",
                            level=3,
                            tables=[tables.scq],
                            condition=lambda: tables.scq.is_available(),
                        ),
                        TableSection(
                            title="Social Responsiveness Scale",
                            level=3,
                            tables=[tables.srs],
                            condition=lambda: tables.srs.is_available(),
                        ),
                    ],
                ),
                PageBreak(),
                ParagraphSection(
                    content="Depression and Anxiety Symptoms",
                    style="Heading 1",
                    condition=lambda: _any_table_available(
                        tables.mfq,
                        tables.scared,
                    ),
                    subsections=[
                        TableSection(
                            title=(
                                "Mood and Feelings Questionnaire (MFQ) - Long Version"
                            ),
                            level=3,
                            tables=[tables.mfq],
                            condition=lambda: tables.mfq.is_available(),
                        ),
                        TableSection(
                            title="Screen for Child Anxiety Related Disorders",
                            level=3,
                            tables=[tables.scared],
                            condition=lambda: tables.scared.is_available(),
                        ),
                    ],
                ),
            ],
        ),
    )


def _any_table_available(*tbls: base.WordTableSection) -> bool:
    """True if any table's data is available, false otherwise."""
    return any(tbl.is_available() for tbl in tbls)


def _all_tables_available(*tbls: base.WordTableSection) -> bool:
    """True if all tables' data are available, false otherwise."""
    return all(tbl.is_available() for tbl in tbls)


def _flatten(collection: Iterable[Iterable[T]]) -> list[T]:
    """Flattens an iterable of iterables into a flat list."""
    return [item for sub_iterable in collection for item in sub_iterable]
