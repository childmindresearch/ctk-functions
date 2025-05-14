"""Utilities for Pyrite reports."""

import abc
import enum
from collections.abc import Callable
from typing import Literal

import cmi_docx
import pydantic
from docx import document
from docx.enum import text as enum_text

from ctk_functions.routers.pyrite.tables import base

VALID_PARAGRAPH_STYLES = Literal[
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 1 Centered",
    "Normal",
    "Normal Hanging",  # After first line, 0.7cm indent.
]


class RunStyles(enum.StrEnum):
    """Valid styles in the Word template for runs."""

    Emphasis = "Emphasis"


class Section(pydantic.BaseModel, abc.ABC):
    """Represents a section in the report structure."""

    subsections: list["Section"] = pydantic.Field(default_factory=list)
    condition: Callable[[], bool] = lambda: True

    def add_to(self, doc: document.Document) -> None:
        """Adds the section to the report."""
        if not self.condition():
            return

        self._add_to(doc)
        for subsection in self.subsections:
            subsection.add_to(doc)

    @abc.abstractmethod
    def _add_to(self, doc: document.Document) -> None:
        """Adds this section to the report.

        Handling of the conditional and subsections is done by self.add_to()
        """


class PageBreak(Section):
    """A page break in the report."""

    def _add_to(self, doc: document.Document) -> None:
        """Adds a page break to the report.

        If any paragraph exists, appends it to the last paragraph. Otherwise,
        create

        Args:
            doc: The document to add the page break to.
        """
        para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
        para.add_run().add_break(enum_text.WD_BREAK.PAGE)


class RunsSection(Section):
    """Represents a text block in the report structure with sub-formatting.

    Attributes:
        content: The textual content of each run.
        run_styles: The style of each run. Can be a valid CharacterStyle,
            a RunStyle object, or None (document's default style).
        condition: Condition to evaluate for section inclusion.
        subsections: Subsections of this section, will be appended if
            this section's condition and the subsection's condition is met.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    content: tuple[str, ...]
    paragraph_style: cmi_docx.ParagraphStyle | VALID_PARAGRAPH_STYLES | None = None
    run_styles: tuple[cmi_docx.RunStyle | RunStyles | None, ...]

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if isinstance(self.paragraph_style, cmi_docx.ParagraphStyle):
            para = doc.add_paragraph("")
            extended_paragraph = cmi_docx.ExtendParagraph(para)
            extended_paragraph.format(self.paragraph_style)
        else:
            para = doc.add_paragraph("", self.paragraph_style)

        for text, style in zip(self.content, self.run_styles, strict=True):
            if style is None or isinstance(style, str):
                para.add_run(text, style)
                continue

            run = para.add_run(text)
            extend_run = cmi_docx.ExtendRun(run)
            extend_run.format(style)


class ParagraphSection(Section):
    """Represents a text block in the report structure.

    Attributes:
        content: The textual content of the paragraph.
        style: Either a ParagraphStyle defining the styling, a string
            matching a Style name in the document, or None. If None,
            uses the document's default style.
        condition: Condition to evaluate for section inclusion.
        subsections: Subsections of this section, will be appended if
            this section's condition and the subsection's condition is met.
    """

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    content: str
    style: cmi_docx.ParagraphStyle | VALID_PARAGRAPH_STYLES | None = None

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if isinstance(self.style, cmi_docx.ParagraphStyle):
            para = doc.add_paragraph(self.content)
            extended_paragraph = cmi_docx.ExtendParagraph(para)
            extended_paragraph.format(self.style)
        else:
            doc.add_paragraph(self.content, self.style)


class TableSection(Section):
    """Represents a table section in the report structure."""

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)

    title: str | None = None
    level: int | None = None
    tables: list[base.WordTableSection] = pydantic.Field(default_factory=list)

    def _add_to(self, doc: document.Document) -> None:
        """Adds the section to a document.

        Args:
            doc: The document to add the section to.
        """
        if self.title and self.level:
            doc.add_heading(self.title, level=self.level)
        for table in self.tables:
            if table.is_available():
                table.add_to(doc)


class ImageSection(Section):
    """Represents a image section in the report structure."""

    model_config = pydantic.ConfigDict(arbitrary_types_allowed=True)
    path: pydantic.FilePath

    def _add_to(self, doc: document.Document) -> None:
        """Adds the image to a document.

        Args:
            doc: The document to add the image to.
        """
        doc.add_picture(str(self.path))
