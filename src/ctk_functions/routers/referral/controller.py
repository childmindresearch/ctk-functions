"""Business logic for the Pyrite endpoints."""

import io

import docx

from ctk_functions.core import config
from ctk_functions.routers.pyrite.tables import base
from ctk_functions.routers.referral import schemas

logger = config.get_logger()
settings = config.get_settings()

DATA_DIR = settings.DATA_DIR


def post_referral(request: schemas.PostReferralRequest) -> bytes:
    """Generates a referral for a given table.

    Args:
        request: The tables to generate a referral for.

    Returns:
        The .docx file bytes.
    """
    doc = docx.Document(str(DATA_DIR / "referral_template.docx"))
    renderers = [_table_to_renderer(table.table) for table in request.tables]
    titles = [table.title for table in request.tables]

    for title, renderer in zip(titles, renderers, strict=True):
        base.ParagraphBlock(content=title, level=2).add_to(doc)
        renderer.add_to(doc)

    out = io.BytesIO()
    doc.save(out)
    doc.save("/Users/reinder.vosdewael/Desktop/test.docx")
    return out.getvalue()


def _table_to_renderer(
    table: dict[str, tuple[str, ...]],
) -> base.WordDocumentTableRenderer:
    """Converts the requested table to a table renderer."""
    headers = [base.WordTableCell(content=key) for key in table]
    row_values = zip(*table.values(), strict=True)
    rows = [[base.WordTableCell(content=text) for text in vals] for vals in row_values]
    markup = base.WordTableMarkup(rows=[headers, *rows])
    return base.WordDocumentTableRenderer(markup=markup)
